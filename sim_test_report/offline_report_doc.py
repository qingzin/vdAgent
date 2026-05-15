

import threading


import numpy as np

import pandas as pd

import matplotlib

from scipy.optimize import curve_fit

from scipy.signal import savgol_filter, butter, filtfilt, medfilt

from scipy.stats import linregress


from control_carsim import ControlCarsim

from utils import (load_configs, configure_matplotlib_style, save_fig_to_word,
                    save_fig_to_cell, safe_idx, make_legend_outside)

from handproc import zxq_proc, jy_proc, wthz_proc, qlzd_proc, qifu_proc, xcxjsd_proc, jmc_proc


matplotlib.use('Agg')

import math

import os
from pathlib import Path

from datetime import datetime

import matplotlib.pyplot as plt

from docx import Document

from docx.shared import Inches, Pt, Cm, RGBColor

from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.enum.table import WD_ALIGN_VERTICAL

from docx.oxml.ns import qn

from docx.oxml import OxmlElement


matplotlib.use('Agg')  # 必须在 pyplot 之前


CONFIG_PATH = str(Path(__file__).resolve().with_name('offline_report_config.json'))

CONFIGS = load_configs(CONFIG_PATH)


DP = 3  # 通常缩写为 d.p.小数位数

G_FS = 100


PAGE_WIDTH = 21.0

PAGE_MARGIN = 1.27

PAGE_USABLE_WIDTH = PAGE_WIDTH - PAGE_MARGIN * 2

IND_WIDTH = 3.0

CAL_WIDTH = 2.5


configure_matplotlib_style()



def find_file_fuzzy(folder_path, keyword):

    """模糊搜索文件夹下的 CSV 文件"""

    if not os.path.exists(folder_path):

        return None

    files = [f for f in os.listdir(folder_path) if keyword in f and f.lower().endswith('.csv')]

    return os.path.join(folder_path, files[0]) if files else None



def safe_read_csv(csv_path):

    """安全读取 CSV，处理编码问题并标准化列名"""

    try:

        try:

            df = pd.read_csv(csv_path, encoding="utf-8", engine="python")

        except:

            df = pd.read_csv(csv_path, encoding="gbk", engine="python")


        if df is None or df.empty:

            print(f"  ⚠️ 文件内容为空: {os.path.basename(csv_path)}")

            return None


        df.columns = [str(c).strip() for c in df.columns]


        raw_alias = CONFIGS.get('common_config', {}).get('col_alias', {})

        alias_map = {}

        for std_name, aliases in raw_alias.items():

            # 标准名本身也要存进去 (小写: 原始标准名)

            alias_map[std_name.lower()] = std_name

            # 所有的别名也指向这个标准名

            for a in aliases:

                alias_map[a.lower()] = std_name


        if alias_map:

            new_columns = {}

            for col in df.columns:

                col_lower = col.lower()

                if col_lower in alias_map:

                    new_columns[col] = alias_map[col_lower]

                else:

                    new_columns[col] = col

            df.rename(columns=new_columns, inplace=True)


        has_string_bugs = False

        for col in df.columns:

            # 检查列中是否包含 str 类型的数据

            str_mask = df[col].apply(lambda x: isinstance(x, str))

            if str_mask.any():

                has_string_bugs = True

                weird_values = df[col][str_mask].unique()[:5]

                print(f"  🔍 [诊断] 列 '{col}' 中混入了字符串数据: {weird_values}")


        if has_string_bugs:

            print("  ⚠️ 数值中存在str，执行数值清洗...剔除无效行")

            df = df.apply(pd.to_numeric, errors='coerce')

            # 处理信号中途偶发的 NaN

            df = df.ffill()

            # 如果最开头就是 NaN 导致无法前向填充

            df = df.fillna(0)

        return df


    except Exception as e:

        print(f"  ❌ 读取异常 [{os.path.basename(csv_path)}]: {e}")

        return None



def set_cell_border(cell, **kwargs):

    """设置单元格边框"""

    tc = cell._tc

    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.first_child_found_in("w:tcBorders")

    if tcBorders is None:

        tcBorders = OxmlElement('w:tcBorders')

        tcPr.append(tcBorders)


    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):

        edge_data = kwargs.get(edge)

        if edge_data:

            tag = 'w:{}'.format(edge)

            element = tcBorders.find(qn(tag))

            if element is None:

                element = OxmlElement(tag)

                tcBorders.append(element)

            for key in ["sz", "val", "color", "space", "shadow"]:

                if key in edge_data:

                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))



def set_run_font(run, size_pt=10.5, bold=False, italic=False, color=None):

    """强制设置中西文双字体 (宋体 + Times New Roman)"""

    run.font.name = 'Times New Roman'

    run.font.size = Pt(size_pt)

    run.font.bold = bold

    run.font.italic = italic

    if color:

        run.font.color.rgb = color


    r = run._element

    rPr = r.get_or_add_rPr()

    fonts = OxmlElement('w:rFonts')

    fonts.set(qn('w:ascii'), 'Times New Roman')

    fonts.set(qn('w:hAnsi'), 'Times New Roman')

    fonts.set(qn('w:eastAsia'), 'SimSun')

    fonts.set(qn('w:cs'), 'Times New Roman')


    old_fonts = rPr.find(qn('w:rFonts'))

    if old_fonts is not None:

        rPr.remove(old_fonts)

    rPr.append(fonts)



class IndicatorCalculator:

    def __init__(self, carsim_controller=None):

        # 结构: {(car, cond): {'metrics': {}, 'figs': {}, 'data': df_proc}}

        self.carsim_controller = carsim_controller

        self._results_cache = {}

        self.crnt_csv_path = None

        self.veh_info_map = {}


    def refresh_csv_path(self, csv_path):

        self.crnt_csv_path = csv_path


    def get_condition_results(self, car_name, cond_key, df, config, car_folder_path=None):

        """核心入口：执行计算并缓存结果、分析图和预处理后的数据"""


        cache_key = (car_name, cond_key)

        if cache_key in self._results_cache:

            return self._results_cache[cache_key]


        if df is None or df.empty:

            return {'metrics': {}, 'figs': {}, 'data': pd.DataFrame()}


        if car_name not in self.veh_info_map:

            if car_folder_path and os.path.exists(car_folder_path):

                txt_path = os.path.join(car_folder_path, 'config.txt')

                print(f"  🔍 [LazyLoad] 正在加载车辆配置: {car_name}")

                self.veh_info_map[car_name] = self.get_veh_config_from_txt(txt_path)

            else:

                # 如果没路径，给个空字典防止后续报错

                self.veh_info_map[car_name] = {}


        common_cfg = config.get('common_config', {})

        cond_cfg = config.get(cond_key, {})


        # 路由到专项计算函数

        ds = cond_cfg.get('Dataset')

        is_multi_condition = "Multi-Condition" in ds



        # 别名转换与基础预处理

        df_proc = self._preprocess_df(df, common_cfg)


        if "Steady-State" in ds:

            df_proc = df_proc[(df_proc['Ay'] > 0.15) & (df_proc['Steer_SW'] > 0.0)]

            res = self._calc_steady_state(df_proc, car_name)

        elif "Pulse Steer" in ds:

            res = self._calc_pulse(df_proc, car_name)

        elif "Central" in ds:

            res = self._calc_center_steer(df_proc, car_name)

        elif "Step" in ds:

            res = self._calc_step(df_proc, car_name)

        # elif "Brake" in ds:

        #     res = self._calc_brake(df_proc, car_name)

        elif "5kmh Brake" in ds:

            res = self._calc_5kmh_brake(df_proc, car_name)

        elif "Low g" in ds:

            res = self._calc_lowg(df_proc, car_name)

        elif "Acceleration Pitch" in ds:

            res = self._calc_acc(df_proc, car_name)

        elif "Braking Pitch" in ds:

            res = self._calc_brake(df_proc, car_name)

        elif "Bump_YIQI" in ds:

            res = self._clac_bump(df_proc, car_name)

        elif is_multi_condition:

            res = self._calc_multi_condition(car_name, car_folder_path, common_cfg)

        else:

            res = self._calc_generic(df_proc, cond_cfg, car_name)

        res['data'] = df_proc

        self._results_cache[cache_key] = res

        return res


    def _preprocess_df(self, df, common_cfg):

        """统一列名映射"""

        alias_map = common_cfg.get('col_alias', {})

        rename_dict = {}

        for std, aliases in alias_map.items():

            for a in aliases:

                if a in df.columns:

                    rename_dict[a] = std

                    break


        df.rename(columns=rename_dict, inplace=True)

        prep_cfg = common_cfg.get('preprocess', {})


        sub_list = prep_cfg.get('sub_ori', [])

        for col in sub_list:

            if col in df.columns:

                df[col] = df[col] - df[col].iloc[0]


        scale_dict = prep_cfg.get('scale', {})


        for col, factor in scale_dict.items():

            if col in df.columns:

                df[col] = df[col] * factor


        return df


    def _calc_center_steer(self, df, car_name):

        """专项：转向响应 - 计算灵敏度并生成标注图"""

        metrics = {}

        figs = {}


        def _fit_hysteresis_loop(x_raw, y_raw):


            try:

                # 1. 预处理

                # win_len = max(5, min(51, len(x_raw) // 10 | 1))

                # x_smooth = savgol_filter(x_raw, window_length=win_len, polyorder=2)

                # dx = np.gradient(x_smooth)

                dx = np.gradient(x_raw)

                threshold = np.max(np.abs(dx)) * 0.1


                idx_rise = np.where(dx > threshold)[0]

                idx_fall = np.where(dx < -threshold)[0]


                if len(idx_rise) < 10 or len(idx_fall) < 10:

                    return None


                # 2. 确定网格范围 (防止外推)

                x_min = max(np.min(x_raw[idx_rise]), np.min(x_raw[idx_fall]))

                x_max = min(np.max(x_raw[idx_rise]), np.max(x_raw[idx_fall]))

                x_fit_grid = np.linspace(x_min, x_max, 200)


                p_rise = np.poly1d(np.polyfit(x_raw[idx_rise], y_raw[idx_rise], 4))

                p_fall = np.poly1d(np.polyfit(x_raw[idx_fall], y_raw[idx_fall], 4))


                # 4. 生成拟合后的 y 值 (m/s^2)

                y_fit_rise = p_rise(x_fit_grid)

                y_fit_fall = p_fall(x_fit_grid)


                # 为了后续求导方便，我们把多项式对象也传出去

                return {

                    'x_grid': x_fit_grid,

                    'y_rise': y_fit_rise,

                    'y_fall': y_fit_fall,

                    'p_rise': p_rise,  # 保存多项式对象

                    'p_fall': p_fall

                }

            except Exception as e:

                return None


        def _calculate_dead_zone(fit_res):

            if not fit_res:

                return None


            x_grid = fit_res['x_grid']

            y_rise = fit_res['y_rise']

            y_fall = fit_res['y_fall']


            def find_zero_crossing_x(x_arr, y_arr):

                # 寻找符号改变的位置

                signs = np.sign(y_arr)

                sign_change_indices = np.where(np.diff(signs))[0]


                if len(sign_change_indices) == 0:

                    # 没有过零点，找绝对值最小的点

                    idx = np.argmin(np.abs(y_arr))

                    return x_arr[idx]


                # 取最靠近中心的过零点 (假设中心在 0 附近)

                candidates = []

                for i in sign_change_indices:

                    x1, x2 = x_arr[i], x_arr[i + 1]

                    y1, y2 = y_arr[i], y_arr[i + 1]

                    if y2 == y1:

                        x_zero = x1

                    else:

                        x_zero = x1 - y1 * (x2 - x1) / (y2 - y1)

                    candidates.append(x_zero)


                # 返回绝对值最小的 X

                return min(candidates, key=abs)


            try:

                x_rise_0 = find_zero_crossing_x(x_grid, y_rise)

                x_fall_0 = find_zero_crossing_x(x_grid, y_fall)

                width = abs(x_rise_0 - x_fall_0)

                return width, x_rise_0, x_fall_0

            except Exception:

                return None


        def _calculate_sensitivity(fit_res):

            if not fit_res:

                return None, None, None, None


            p_rise = fit_res['p_rise']

            p_fall = fit_res['p_fall']


            # 获取原始网格数据

            x_grid = fit_res['x_grid']

            y_rise = fit_res['y_rise']

            y_fall = fit_res['y_fall']


            # --- 内部辅助函数：计算单条曲线的指标和坐标 ---

            def analyze_curve(poly_obj, y_g_vals):

                raw_slopes = poly_obj.deriv()(x_grid)


                # 2. 寻找最小灵敏度点 (中心区)

                mask_center = (y_g_vals < 0.1) & (y_g_vals > 0)

                if np.any(mask_center):

                    center_indices = np.where(mask_center)[0]

                    idx_min = center_indices[np.argmin(np.abs(raw_slopes[center_indices]))]

                else:

                    idx_min = np.argmin(raw_slopes)


                idx_01g = np.argmin(np.abs(y_g_vals - 0.1))


                m_sens_metric = raw_slopes[idx_min] * 100.0

                s_01g_metric = raw_slopes[idx_01g] * 100.0


                viz_min = {

                    'x': x_grid[idx_min],

                    'y': y_g_vals[idx_min],

                    'slope': raw_slopes[idx_min]

                }


                viz_01g = {

                    'x': x_grid[idx_01g],

                    'y': y_g_vals[idx_01g],

                    'slope': raw_slopes[idx_01g]

                }


                return m_sens_metric, s_01g_metric, viz_min, viz_01g


            min_rise, s01g_rise, v_min_rise, v_01g_rise = analyze_curve(p_rise, y_rise)

            min_fall, s01g_fall, v_min_fall, v_01g_fall = analyze_curve(p_fall, y_fall)


            # 取平均

            final_min_sens = (min_rise + min_fall) / 2.0

            print(min_rise, min_fall)

            final_sens_01g = (s01g_rise + s01g_fall) / 2.0

            print(s01g_rise, s01g_fall)


            viz_rise = (v_min_rise, v_01g_rise)

            viz_fall = (v_min_fall, v_01g_fall)


            return final_min_sens, final_sens_01g, viz_rise, viz_fall


        def calc_vda():

            result_container = {

                # "v_dz": 0,

                "v_min": 0,

                "v_01g": 0,

                "fig": None

            }


            def target_task():

                temp_filename = ""

                try:

                    temp_filename = os.path.abspath(f"temp_vda_{car_name}_{datetime.now().strftime('%H%M%S%f')}.csv")

                    target_cols = ["TimeStep", "Steer_SW", "Ay", "M_SW", "Avz"]

                    existing_cols = [col for col in target_cols if col in df.columns]


                    df_temp = df[existing_cols].copy()

                    if 'Ay' in df_temp.columns:

                        df_temp['Ay'] = df_temp['Ay'] * 9.8

                    df_temp.to_csv(temp_filename, index=False)

                    cols = list(df_temp.columns)


                    zxq_channel_list = [

                        safe_idx("TimeStep", cols),

                        safe_idx("Steer_SW", cols),

                        safe_idx("Ay", cols),

                        safe_idx("M_SW", cols),

                        safe_idx("Avz", cols),

                        safe_idx("Avz", cols),

                    ]


                    # vda

                    res_dict = zxq_proc([temp_filename], zxq_channel_list)


                    if res_dict and '平均值' in res_dict and len(res_dict['平均值']) > 0:

                        avg_data = res_dict['平均值']

                        result_container['v_min'] = round(avg_data[0].get('11-最小灵敏度(g/100deg)', 0), DP)

                        result_container['v_01g'] = round(avg_data[0].get('12-0.1g灵敏度(g/100deg)', 0), DP)

                        # result_container['v_dz'] = round(avg_data[0].get('17-转向力矩转角迟滞(deg)', 0), DP)


                        if len(avg_data) > 1 and isinstance(avg_data[1], dict):

                            plots_dict = avg_data[1]

                            target_key = '侧向加速度-方向盘转角'


                            if target_key in plots_dict:

                                plot_data = plots_dict[target_key]


                                fig_vda, ax_v = plt.subplots()


                                # 遍历 curve 列表画图

                                if 'curve' in plot_data:

                                    for curve_info in plot_data['curve']:

                                        x_data = curve_info.get('x')

                                        y_data = curve_info.get('y')

                                        name = curve_info.get('curvename', 'Curve')

                                        ctype = curve_info.get('curvetype', 'line')


                                        # 确保数据有效

                                        if x_data is not None and y_data is not None and len(x_data) > 0:

                                            if ctype == 'scatter':

                                                ax_v.scatter(x_data, y_data, s=4, alpha=0.2, color='gray', label=name)

                                            else:

                                                # 区分左转/右转颜色

                                                c = 'red' if '左' in name else ('blue' if '右' in name else 'green')

                                                ax_v.plot(x_data, y_data, lw=2, color=c, label=name)


                                ax_v.set_title(f"[VDA] {car_name} 中心区")

                                ax_v.set_xlabel(plot_data.get('xlabel', 'Steer [deg]'))

                                ax_v.set_ylabel(plot_data.get('ylabel', 'Ay [m/s^2]'))

                                make_legend_outside()

                                result_container['fig'] = fig_vda

                    else:

                        print(f"\t[VDA] {car_name}计算返回结果为空")


                except Exception as e:

                    print(f"\t[VDA] 线程内计算出错: {e}")

                finally:

                    if temp_filename and os.path.exists(temp_filename):

                        try:

                            os.remove(temp_filename)

                        except:

                            pass


            # 线程

            t = threading.Thread(target=target_task)

            t.daemon = True

            t.start()


            t.join(timeout=50)


            if t.is_alive():

                print(f"[VDA] 警告: 计算超时 (超过50s)，已跳过！")


            return tuple(result_container.values())

        def calc_lb():

            x_steer = df['Steer_SW'].values

            y_ay = df['Ay'].values  # 单位 g


            p_steer_vs_ay = np.poly1d(np.polyfit(y_ay, x_steer, 4))


            p_ay_vs_steer = np.poly1d(np.polyfit(x_steer, y_ay, 5))

            dp5 = p_ay_vs_steer.deriv()


            ay_targets = np.linspace(-0.1, 0.1, 21)

            steer_at_low_ay = p_steer_vs_ay(ay_targets)


            rates = dp5(steer_at_low_ay)


            val3_min = round(np.min(np.abs(rates)) * 100, DP)

            val3_01g = round(np.abs(rates[-1]) * 100, DP)


            return val3_min, val3_01g


        def calc_my():


            dz_res, val1_min, val1_01g = 0, 0, 0

            fig_py = None

            width = 0


            if 'Steer_SW' in df.columns and 'Ay' in df.columns:

                fit_res = _fit_hysteresis_loop(df["Steer_SW"], df["Ay"])


                if fit_res:

                    # B. 计算死区

                    dz_res = _calculate_dead_zone(fit_res)


                    sens_res = _calculate_sensitivity(fit_res)

                    viz_rise, viz_fall = None, None


                    if sens_res[0] is not None:

                        val1_min = round(sens_res[0], DP)

                        val1_01g = round(sens_res[1], DP)

                        # 获取用于可视化的坐标点数据

                        viz_rise = sens_res[2]  # (min_point_dict, 01g_point_dict)

                        viz_fall = sens_res[3]


                    fig_py, ax = plt.subplots()


                    ax.plot(df['Steer_SW'], df['Ay'], color='gray', alpha=0.3, label='Raw Data', zorder=1)

                    ax.plot(fit_res['x_grid'], fit_res['y_rise'], 'r-', lw=1, alpha=0.5, label='Rise Fit', zorder=5)

                    ax.plot(fit_res['x_grid'], fit_res['y_fall'], 'b-', lw=1, alpha=0.5, label='Fall Fit', zorder=5)

                    if dz_res:

                        width, x1, x2 = dz_res

                        ax.axvline(x=x1, color='k', linestyle='--', alpha=0.3, lw=0.5)

                        ax.axvline(x=x2, color='k', linestyle='--', alpha=0.3, lw=0.5)

                        ax.axvspan(min(x1, x2), max(x1, x2), color='yellow', alpha=0.1, label='Deadband')

                        mid_x = (x1 + x2) / 2

                        ax.text(mid_x, 0, f'{width:.1f}deg', ha='center', va='center',

                                fontsize=8, bbox=dict(facecolor='white', alpha=0.8, edgecolor='none'))


                    def plot_tangent(ax, point_info, color, label_prefix):

                        if not point_info: return

                        x0 = point_info['x']

                        y0 = point_info['y']

                        slope = point_info['slope']  # 注意：这里的 slope 是 dy/dx (g/deg)


                        # 画点

                        ax.plot(x0, y0, marker='*', color=color, zorder=10,

                                label=f'{label_prefix} Loc')


                        # 画一小段切线 (长度 2 deg)

                        dx = 4.0

                        x_line = np.array([x0 - dx, x0 + dx])

                        y_line = slope * (x_line - x0) + y0

                        ax.plot(x_line, y_line, linestyle='-', color=color, lw=2, alpha=0.6)


                    if viz_rise and viz_fall:

                        # viz_rise[0] 是最小灵敏度点, viz_rise[1] 是 0.1g 点


                        # 绘制最小灵敏度点 (通常在 0g 附近) -> 用青色 (Cyan)

                        plot_tangent(ax, viz_rise[0], 'cyan', 'Min Sens')

                        plot_tangent(ax, viz_fall[0], 'cyan', '')  # Fall曲线的点不重复加图例


                        # 绘制 0.1g 灵敏度点 -> 用洋红色 (Magenta)

                        plot_tangent(ax, viz_rise[1], 'magenta', '0.1g Sens')

                        plot_tangent(ax, viz_fall[1], 'magenta', '')


                    # 装饰图表

                    ax.set_title(f"[方法2] {car_name} Python PolyFit\nDeadband={width:.2f}deg | MinSens={val1_min}")

                    ax.set_xlabel('Steer Angle [deg]')

                    ax.set_ylabel('Ay [g]')

                    make_legend_outside()


            return round(dz_res[0], DP), round(val1_min, DP), round(val1_01g, DP), fig_py


        my_dz, my_min, my_01g, fig_py = calc_my()

        vda_min, vda_01g, fig_vda = calc_vda()

        # vda_dz, vda_min, vda_01g, fig_vda = calc_vda()

        val3_min, val3_01g = calc_lb()


        metrics['转向死区(deg)'] = [my_dz]

        figs['转向死区(deg)'] = fig_py

        metrics['最小转向灵敏度(g/100deg)'] = [vda_min, my_min, val3_min]

        metrics['0.1g转向灵敏度(g/100deg)'] = [vda_01g, my_01g, val3_01g]


        combined_figs = []

        if fig_py: combined_figs.append(fig_py)

        if fig_vda: combined_figs.append(fig_vda)


        if combined_figs:

            figs['最小转向灵敏度(g/100deg)'] = combined_figs

            figs['0.1g转向灵敏度(g/100deg)'] = combined_figs  # 两个指标共用同一组图


        return {'metrics': metrics, 'figs': figs}


    def _calc_steady_state(self, df, car_name):

        """专项：稳态回转 - 计算梯度并生成拟合分析图"""

        metrics = {}

        figs = {}


        if self.carsim_controller is None:

            return {'metrics': metrics, 'figs': figs}


        # 优先级是：config.txt > CSV 数据 > 默认值 1.0。

        steer_ratio = 1.0

        crnt_veh_f_wheel_base = 1.0

        crnt_veh_r_wheel_base = 1.0

        if "Steer_SW" and "Steer_L1" in df.columns:

            steer_ratio = df["Steer_SW"].mean() / df["Steer_L1"].mean()

            print(f"\t[{car_name}] 从 CSV 提取传动比: {steer_ratio:.2f}")

        else:

            print(f"\t[{car_name}] ⚠️ CSV未找到传动比，使用config.txt")

            # 无论 CSV 里有没有传动比，先从 txt (veh_info_map) 加载轴距等参数


        veh_params = self.veh_info_map.get(car_name, {})


        if veh_params:

            if 5.0 <= veh_params.get('steer_ratio', 0) <= 20:

                steer_ratio = veh_params['steer_ratio']


            else:

                print(f"\t[{car_name}] 从 config.txt 传动比: {veh_params['steer_ratio']:.2f} 数值异常，使用 {steer_ratio} 替代计算")


            if veh_params.get('f_wheel_base', 0) > 0:

                crnt_veh_f_wheel_base = veh_params['f_wheel_base']

                crnt_veh_r_wheel_base = veh_params['r_wheel_base']

                # 记录到 metrics 以便报告输出 (统一写成列表格式 [val]，对齐单方法)

                metrics['质心到前轴距离(m)'] = [crnt_veh_f_wheel_base]

                metrics['质心到后轴距离(m)'] = [crnt_veh_r_wheel_base]


        metrics['传动比(-)'] = [round(steer_ratio, 2)]


        print(

            f"\t[{car_name}] 当前参数 -> 传动比:{steer_ratio:.2f}, 轴距:前{crnt_veh_f_wheel_base}/后{crnt_veh_r_wheel_base}")

        def roll_gradient(df, car_name=""):

            mask = (df["Ay"] >= 0.1) & (df["Ay"] <= 0.65)


            x_fit = df.loc[mask, "Ay"].values

            y_fit = df.loc[mask, "Roll"].values


            if len(x_fit) < 5 or df["Ay"].max() < 0.35:

                print(f"\t[{car_name}] 侧倾梯度：0.1g-0.65g 数据点不足，无法拟合")

                return 0, None


            p_coeffs = np.polyfit(x_fit, y_fit, 1)

            k = p_coeffs[0]


            fig_roll, ax_roll = plt.subplots()

            ax_roll.scatter(df["Ay"], df["Roll"], color='gray', s=3, alpha=0.15, label='Raw Data (All)')

            ax_roll.scatter(x_fit, y_fit, color='black', s=6, alpha=0.5, label='Fit Range (0.1-0.5g)')


            x_plot = np.linspace(0.1, 0.65, 50)

            y_plot = np.polyval(p_coeffs, x_plot)

            ax_roll.plot(x_plot, y_plot, 'r-', linewidth=2.5, label=f'Linear Fit (K={k:.2f})')


            ax_roll.set_title(f'[方法2] {car_name} 侧倾梯度\nK = {k:.3f} deg/g')

            ax_roll.set_xlabel('Ay [g]')

            ax_roll.set_ylabel('Roll [deg]')


            ax_roll.set_xlim(0, max(df["Ay"].max() * 1.1, 0.6))

            ax_roll.set_ylim(0, max(df["Roll"].max() * 1.1, y_plot.max() + 0.5))

            make_legend_outside(ax_roll)


            return round(k, DP), fig_roll


        def calc_vda0():

            result_container = {

                "v_03g": 0,

                "v_05g": 0,

                "v_07g": 0,

                'roll_gradient': 0,

                "f_cepian": 0,

                "r_cepian": 0,

                "v_max_ay":0,

                "fig": None

            }


            def target_task():

                temp_filename = ""

                try:

                    temp_filename = os.path.abspath(f"temp_vda_{car_name}_{datetime.now().strftime('%H%M%S%f')}.csv")

                    target_cols = ["TimeStep", "Steer_SW", "Ay", "Avz", "Roll", "M_SW", "Beta","Vx"]

                    existing_cols = [col for col in target_cols if col in df.columns]


                    df_temp = df[existing_cols].copy()

                    if 'Ay' in df_temp.columns:

                        df_temp['Ay'] = df_temp['Ay'] * 9.8

                    df_temp.to_csv(temp_filename, index=False)

                    cols = list(df_temp.columns)


                    wthz_channel_list = [

                        safe_idx("TimeStep", cols),

                        safe_idx("Steer_SW", cols),

                        safe_idx("Ay", cols),

                        safe_idx("Avz", cols),

                        safe_idx("Roll", cols),

                        safe_idx("M_SW", cols),

                        safe_idx("Beta", cols),

                        safe_idx("Vx", cols)

                    ]


                    # vda

                    safe_ratio = steer_ratio if (steer_ratio and steer_ratio > 0) else 1.0

                    res_dict = wthz_proc([temp_filename], wthz_channel_list, safe_ratio,crnt_veh_f_wheel_base*1000,crnt_veh_r_wheel_base*1000,0)

                    if res_dict and '平均值' in res_dict and len(res_dict['平均值']) > 0:

                        avg_data = res_dict['平均值']

                        result_container['v_03g'] = round(avg_data[0].get('12-0.3g不足转向度(deg/g)', 0), DP)

                        result_container['v_05g'] = round(avg_data[0].get('13-0.5g不足转向度(deg/g)', 0), DP)

                        result_container['v_07g'] = round(avg_data[0].get('14-0.7g不足转向度(deg/g)', 0), DP)

                        result_container['roll_gradient'] = round(avg_data[0].get('18-车身侧倾梯度(deg/g)', 0), DP)

                        result_container['f_cepian'] = abs(round(avg_data[0].get('27-0.3g前轴等效侧偏柔度(deg/g)', 0), DP))

                        result_container['r_cepian'] = abs(round(avg_data[0].get('28-0.3g后轴等效侧偏柔度(deg/g)', 0), DP))

                        result_container['v_max_ay'] = round(avg_data[0].get('19-最大侧向加速度(g)', 0), DP)


                        if len(avg_data) > 1 and isinstance(avg_data[1], dict):

                            plots_dict = avg_data[1]

                            target_key = '方向盘转角-侧向加速度'


                            if target_key in plots_dict:

                                plot_data = plots_dict[target_key]


                                fig_vda, ax_v = plt.subplots()

                                # 遍历 curve 列表画图

                                if 'curve' in plot_data:

                                    for curve_info in plot_data['curve']:

                                        x_data = curve_info.get('x')

                                        y_data = curve_info.get('y')

                                        name = curve_info.get('curvename', 'Curve')

                                        ctype = curve_info.get('curvetype', 'line')


                                        # 确保数据有效

                                        if x_data is not None and y_data is not None and len(x_data) > 0:

                                            # 1. 如果是原始数据 -> 灰色半透明散点，作为背景

                                            if '原始' in name:

                                                ax_v.scatter(x_data, y_data, s=3, alpha=0.2, color='gray', label=name,

                                                             zorder=1)


                                            # 2. 如果是拟合数据 -> 彩色加粗实线，区分左右转

                                            elif '拟合' in name:

                                                c = 'red' if '左' in name else ('blue' if '右' in name else 'green')

                                                ax_v.plot(x_data, y_data, lw=2.5, color=c, label=name, zorder=5)


                                            # 3. 兜底或其他类型

                                            else:

                                                if ctype == 'scatter':

                                                    ax_v.scatter(x_data, y_data, s=4, alpha=0.5, label=name)

                                                else:

                                                    ax_v.plot(x_data, y_data, lw=1.5, linestyle='--', label=name)


                                ax_v.set_title(f"[VDA] {car_name} Ay vs Steer")

                                ax_v.set_xlabel(plot_data.get('xlabel', 'Steer [deg]'))

                                ax_v.set_ylabel(plot_data.get('ylabel', 'Ay [m/s^2]'))

                                make_legend_outside()

                                result_container['fig'] = fig_vda

                    else:

                        print(f"[VDA-3次] {car_name}计算返回结果为空")


                except Exception as e:

                    print(f"[VDA] 线程内计算出错: {e}")

                finally:

                    if temp_filename and os.path.exists(temp_filename):

                        try:

                            os.remove(temp_filename)

                        except:

                            pass


            # 线程

            t = threading.Thread(target=target_task)

            t.daemon = True

            t.start()


            t.join(timeout=50)


            if t.is_alive():

                print(f"[VDA] 警告: 计算超时 (超过50s)，已跳过！")


            return tuple(result_container.values())


        def calc_vda1():

            result_container = {

                "v_03g": 0,

                "v_05g": 0,

                "v_07g": 0,

                'roll_gradient': 0,

                "f_cepian": 0,

                "r_cepian": 0,

                "fig": None

            }


            def target_task():

                temp_filename = ""

                try:

                    temp_filename = os.path.abspath(f"temp_vda_{car_name}_{datetime.now().strftime('%H%M%S%f')}.csv")

                    target_cols = ["TimeStep", "Steer_SW", "Ay", "Avz", "Roll", "M_SW", "Beta","Vx"]

                    existing_cols = [col for col in target_cols if col in df.columns]


                    df_temp = df[existing_cols].copy()

                    if 'Ay' in df_temp.columns:

                        df_temp['Ay'] = df_temp['Ay'] * 9.8

                    df_temp.to_csv(temp_filename, index=False)

                    cols = list(df_temp.columns)


                    wthz_channel_list = [

                        safe_idx("TimeStep", cols),

                        safe_idx("Steer_SW", cols),

                        safe_idx("Ay", cols),

                        safe_idx("Avz", cols),

                        safe_idx("Roll", cols),

                        safe_idx("M_SW", cols),

                        safe_idx("Beta", cols),

                        safe_idx("Vx", cols)

                    ]


                    # vda

                    safe_ratio = steer_ratio if (steer_ratio and steer_ratio > 0) else 1.0

                    res_dict = wthz_proc([temp_filename], wthz_channel_list, safe_ratio,crnt_veh_f_wheel_base*1000,crnt_veh_r_wheel_base*1000,1)

                    if res_dict and '平均值' in res_dict and len(res_dict['平均值']) > 0:

                        avg_data = res_dict['平均值']

                        result_container['v_03g'] = round(avg_data[0].get('12-0.3g不足转向度(deg/g)', 0), DP)

                        result_container['v_05g'] = round(avg_data[0].get('13-0.5g不足转向度(deg/g)', 0), DP)

                        result_container['v_07g'] = round(avg_data[0].get('14-0.7g不足转向度(deg/g)', 0), DP)

                        result_container['roll_gradient'] = round(avg_data[0].get('18-车身侧倾梯度(deg/g)', 0), DP)

                        result_container['f_cepian'] = abs(round(avg_data[0].get('27-0.3g前轴等效侧偏柔度(deg/g)', 0), DP))

                        result_container['r_cepian'] = abs(round(avg_data[0].get('28-0.3g后轴等效侧偏柔度(deg/g)', 0), DP))


                        if len(avg_data) > 1 and isinstance(avg_data[1], dict):

                            plots_dict = avg_data[1]

                            target_key = '方向盘转角-侧向加速度'


                            if target_key in plots_dict:

                                plot_data = plots_dict[target_key]


                                fig_vda, ax_v = plt.subplots()

                                # 遍历 curve 列表画图

                                if 'curve' in plot_data:

                                    for curve_info in plot_data['curve']:

                                        x_data = curve_info.get('x')

                                        y_data = curve_info.get('y')

                                        name = curve_info.get('curvename', 'Curve')

                                        ctype = curve_info.get('curvetype', 'line')


                                        # 确保数据有效

                                        if x_data is not None and y_data is not None and len(x_data) > 0:

                                            # 1. 如果是原始数据 -> 灰色半透明散点，作为背景

                                            if '原始' in name:

                                                ax_v.scatter(x_data, y_data, s=3, alpha=0.2, color='gray', label=name,

                                                             zorder=1)


                                            # 2. 如果是拟合数据 -> 彩色加粗实线，区分左右转

                                            elif '拟合' in name:

                                                c = 'red' if '左' in name else ('blue' if '右' in name else 'green')

                                                ax_v.plot(x_data, y_data, lw=2.5, color=c, label=name, zorder=5)


                                            # 3. 兜底或其他类型

                                            else:

                                                if ctype == 'scatter':

                                                    ax_v.scatter(x_data, y_data, s=4, alpha=0.5, label=name)

                                                else:

                                                    ax_v.plot(x_data, y_data, lw=1.5, linestyle='--', label=name)


                                ax_v.set_title(f"[VDA-多次] {car_name} Ay vs Steer")

                                ax_v.set_xlabel(plot_data.get('xlabel', 'Steer [deg]'))

                                ax_v.set_ylabel(plot_data.get('ylabel', 'Ay [m/s^2]'))

                                make_legend_outside()

                                result_container['fig'] = fig_vda

                    else:

                        print(f"[VDA] {car_name}计算返回结果为空")


                except Exception as e:

                    print(f"[VDA] 线程内计算出错: {e}")

                finally:

                    if temp_filename and os.path.exists(temp_filename):

                        try:

                            os.remove(temp_filename)

                        except:

                            pass


            # 线程

            t = threading.Thread(target=target_task)

            t.daemon = True

            t.start()


            t.join(timeout=50)


            if t.is_alive():

                print(f"[VDA] 警告: 计算超时 (超过50s)，已跳过！")


            return tuple(result_container.values())

        def calc_my(df, ay_need):

            # 0. ------ 预处理 ------

            valid_mask = (df["TimeStep"] >= 10) & (df["Ay"] > 0.05)


            # 数据太少保护

            if np.sum(valid_mask) < 20:

                return 0, {}


            # 1. ------ 数据截取 (只取到峰值) ------

            idx_start = df[valid_mask].index[0]


            try:

                idx_max_ay = df.loc[idx_start:, "Ay"].idxmax()

            except:

                idx_max_ay = df.index[-1]


            df_fit = df.loc[idx_start:idx_max_ay].copy()


            # 排序并去重

            df_fit.sort_values(by="Ay", inplace=True)

            df_fit.drop_duplicates(subset=["Ay"], inplace=True)


            x_vals = df_fit["Ay"].values

            y_raw = df_fit["Steer_SW"].values


            if len(x_vals) < 10: return 0, {}


            y_start = y_raw[0]

            x_start = x_vals[0]


            y_fit_data = y_raw - y_start

            x_fit_data = x_vals - x_start


            def power_model(x, c, a, n):

                return c * x + a * np.power(x, n)


            try:

                # bounds 参数说明:

                # c >= 0 (线性刚度非负)

                # a >= 0 (非线性增益非负)

                # n >= 3 (指数至少是3，最高给到 15，防止过拟合)

                popt, _ = curve_fit(power_model, x_fit_data, y_fit_data,

                                    bounds=([0, 0, 3], [np.inf, np.inf, 15]))


                c, a, n = popt

                # print(f"\t[AOS]{ay_need}g 自动拟合出的最佳指数 n = {n:.2f}")


            except Exception as e:

                print(f"拟合失败: {e}")

                return 0, {}


            # 4. ------ 计算目标点指标 ------

            x_target_shifted = ay_need - x_start


            if x_target_shifted < 0:

                k_sw = c

                target_y_shifted = 0

            else:

                k_sw = c + a * n * (x_target_shifted ** (n - 1))

                target_y_shifted = power_model(x_target_shifted, c, a, n)


            target_y = target_y_shifted + y_start

            k_rw = k_sw / steer_ratio


            # 5. ------ 可视化 ------

            fig1, ax1 = plt.subplots()


            # 画原始点

            ax1.scatter(df['Ay'], df['Steer_SW'], color='gray', s=3, alpha=0.2, label='Raw Data')

            x_plot = np.linspace(x_vals.min(), x_vals.max(), 100)

            y_plot = power_model(x_plot - x_start, c, a, n) + y_start

            ax1.plot(x_plot, y_plot, 'r-', linewidth=2.5, label='c * x + a * np.power(x, n)')

            ax1.plot(ay_need, target_y, 'go', markersize=8, zorder=10)

            tangent_len = 0.15

            x_tan = np.array([ay_need - tangent_len, ay_need + tangent_len])

            y_tan = k_sw * (x_tan - ay_need) + target_y

            ax1.plot(x_tan, y_tan, 'g--', linewidth=1.5, label=f'K={k_sw:.1f}')

            ax1.set_title(f'[方法3]{car_name} @{ay_need}g\nK={k_rw:.3f} deg/g')

            ax1.set_xlabel('Ay [g]')

            ax1.set_ylabel('Steer_SW [deg]')

            ax1.set_xlim(0.05, 0.75)

            idx_05 = (df['Ay'] - 0.05).abs().idxmin()

            idx_75 = (df['Ay'] - 0.75).abs().idxmin()

            steer_at_05 = df.loc[idx_05, 'Steer_SW']

            steer_at_75 = df.loc[idx_75, 'Steer_SW']

            ax1.set_ylim(steer_at_05, steer_at_75)

            make_legend_outside()


            return round(k_rw, DP), fig1


        def calc_lb(df):

            filter_time = (df['TimeStep'] >= 10)

            Ay = df['Ay']


            def lb_under(ay_need):

                filter = filter_time & (Ay >= ay_need - 0.05) & (Ay <= ay_need + 0.05)

                x_vals = Ay[filter].values

                y_vals = df['Steer_SW'][filter].values


                if len(Ay[filter]) == 0 or len(df['Steer_SW'][filter]) == 0:

                    print(f"\t[LB方法不足转向度] 数据点不足，无法拟合 {ay_need}g 处梯度")

                    return 0, {}

                p_coeffs = np.polyfit(Ay[filter], df['Steer_SW'][filter], 4)

                p_func = np.poly1d(p_coeffs)

                p_deriv = np.polyder(p_coeffs)

                k_sw = np.polyval(p_deriv, ay_need)

                buzu = k_sw / steer_ratio


                fig_under, ax_under = plt.subplots()

                ax_under.scatter(df['Ay'], df['Steer_SW'], color='gray', s=3, alpha=0.2, label='Raw Data')


                x_smooth = np.linspace(x_vals.min(), x_vals.max(), 100)

                y_smooth = p_func(x_smooth)

                target_y = p_func(ay_need)

                ax_under.plot(x_smooth, y_smooth, 'r-', linewidth=2.5, label='LB PolyFit (order=4)')


                ax_under.plot(ay_need, target_y, 'go', markersize=8, zorder=10)

                tan_len = 0.08  # 切线长度

                x_tan = np.array([ay_need - tan_len, ay_need + tan_len])

                y_tan = k_sw * (x_tan - ay_need) + target_y

                ax_under.plot(x_tan, y_tan, 'g--', linewidth=1.5, label=f'K={k_sw:.1f}')


                ax_under.set_title(f'[方法4]{car_name} @{ay_need}g\nK={buzu:.3f} deg/g')

                ax_under.set_xlabel('Ay [g]')

                ax_under.set_ylabel('Steer_SW [deg]')


                ax_under.set_xlim(0.05, 0.75)

                idx_05 = (df['Ay'] - 0.05).abs().idxmin()

                idx_75 = (df['Ay'] - 0.75).abs().idxmin()

                steer_at_05 = df.loc[idx_05, 'Steer_SW']

                steer_at_75 = df.loc[idx_75, 'Steer_SW']

                ax_under.set_ylim(steer_at_05, steer_at_75)

                make_legend_outside()


                return buzu, fig_under


            # --- 2. 数据有效性检查 ---


            buzu_3g, buzu_3g_fig = lb_under(0.3)

            print(f"\t[方法4]0.3g 不足转向度: {buzu_3g:.3f} deg/g")

            buzu_5g, buzu_5g_fig = lb_under(0.5)

            print(f"\t[方法4]0.5g 不足转向度: {buzu_5g:.3f} deg/g")

            buzu_7g, buzu_7g_fig = lb_under(0.7)

            print(f"\t[方法4]0.7g 不足转向度: {buzu_7g:.3f} deg/g")


            def cepian():

                # --- 4. 前后轴等效侧偏柔度 (Compliance) ---

                # --- 计算侧偏角(SlipAngles) ---

                # 利用二自由度模型公式计算前后轴等效侧偏角

                # Vx需要从km / h转换为m / s(代码中使用了 / 3.6)

                beta_front = (df['Beta'] + crnt_veh_f_wheel_base * df['Avz'] /
                              (df['Vx'] / 3.6) - df['Steer_SW'] / steer_ratio)

                beta_rear = df['Beta'] - crnt_veh_r_wheel_base * df['Avz'] / (df['Vx'] / 3.6)


                filter_03g = filter_time & (Ay >= 0.3 - 0.02) & (Ay <= 0.3 + 0.02)


                if len(Ay[filter_03g]) == 0 or len(df['Steer_SW'][filter_03g]) == 0:

                    print(f"\t[方法2 等效侧偏柔度] 数据点不足，无法拟合")

                    return 0, 0


                p_front = np.polyfit(Ay[filter_03g], beta_front[filter_03g], 3)  # 拟合 Ay 与 侧偏角

                d_front = np.polyder(p_front)  # 求导

                front_cepian_roudu = abs(np.polyval(d_front, 0.3))  # 求 0.3g 处的斜率

                print(f"\t[方法2]前轴等效侧偏柔度：{front_cepian_roudu:.3f} [deg/g]")


                # 后轴

                p_rear = np.polyfit(Ay[filter_03g], beta_rear[filter_03g], 3)

                d_rear = np.polyder(p_rear)

                rear_cepian_roudu = abs(np.polyval(d_rear, 0.3))

                print(f"\t[方法2]后轴等效侧偏柔度：{rear_cepian_roudu:.3f} [deg/g]")


                return front_cepian_roudu, rear_cepian_roudu


            front_cepian_roudu, rear_cepian_roudu = cepian()


            return round(buzu_3g, DP), round(buzu_5g, DP), round(buzu_7g, DP), round(front_cepian_roudu, DP), round(

                rear_cepian_roudu, DP), buzu_3g_fig, buzu_5g_fig, buzu_7g_fig


        my_roll, figs['0.3g侧倾梯度(deg/g)'] = roll_gradient(df)


        my_03g, my_03g_fig = calc_my(df, 0.3)

        my_05g, my_05g_fig = calc_my(df, 0.5)

        my_07g, my_07g_fig = calc_my(df, 0.7)


        lb_03g, lb_05g, lb_07g, lb_f_cepian, lb_r_cepian, lb_03g_fig, lb_05g_fig, lb_07g_fig = calc_lb(

            df)

        vda0_03g, vda0_05g, vda0_07g, vda0_roll, vda0_f_cepian, vda0_r_cepian, v_max_ay, vda0_fig = calc_vda0()


        vda1_03g, vda1_05g, vda1_07g, vda1_roll, vda1_f_cepian, vda1_r_cepian, vda1_fig = calc_vda1()


        metrics['0.3g侧倾梯度(deg/g)'] = [vda0_roll, my_roll]

        metrics['0.3g不足转向度(deg/g)'] = [vda0_03g, vda1_03g, my_03g,  lb_03g]

        figs['0.3g不足转向度(deg/g)'] = [vda0_fig, vda1_fig, my_03g_fig, lb_03g_fig]

        metrics['0.5g不足转向度(deg/g)'] = [vda0_05g, vda1_05g, my_05g, lb_05g]

        figs['0.5g不足转向度(deg/g)'] = [vda0_fig,vda1_fig,  my_05g_fig, lb_05g_fig]

        metrics['0.7g不足转向度(deg/g)'] = [vda0_07g, vda1_07g, my_07g, lb_07g]

        figs['0.7g不足转向度(deg/g)'] = [vda0_fig, vda1_fig, my_07g_fig,lb_07g_fig]

        metrics['最大侧向加速度(g)'] = [v_max_ay]


        metrics['前轴等效侧偏柔度(deg/g)'] = [vda0_f_cepian]

        metrics['后轴等效侧偏柔度(deg/g)'] = [vda0_r_cepian]


        return {'metrics': metrics, 'figs': figs}


    def _calc_lowg(self, df, car_name):

        metrics = {}

        figs = {}


        def calc_my():

            ay_min = 0.4 / 9.81

            ay_max = 1.3 / 9.81

            mask = (df["Ay"] >= ay_min) & (df["Ay"] <= ay_max)

            x_fit = df.loc[mask, "Steer_SW"].values

            y_fit = df.loc[mask, "Ay"].values


            if len(x_fit) < 5:

                print(f"\t[方法2] 线性区灵敏度：0.4m/s2-1.3m/s2 数据点不足，无法拟合")

                return 0, None


            p_coeffs = np.polyfit(x_fit, y_fit, 1)

            k = p_coeffs[0]*100


            fig_lowg, ax_lowg = plt.subplots()

            ax_lowg.scatter(df["Steer_SW"], df["Ay"], color='gray', s=3, alpha=0.15, label='Raw Data (All)')

            ax_lowg.scatter(x_fit, y_fit, color='black', s=6, alpha=0.5, label='Fit Range (0.4m/s2-1.3m/s2)')


            x_plot = np.linspace(x_fit.min(), x_fit.max(), 50)

            y_plot = np.polyval(p_coeffs, x_plot)

            ax_lowg.plot(x_plot, y_plot, 'r-', linewidth=2.5, label=f'Linear Fit (K={k:.2f})')


            ax_lowg.set_title(f'[方法2] {car_name} 线性区灵敏度\n K = {k:.3f} g/100°')

            ax_lowg.set_xlabel('Steer_SW [deg]')

            ax_lowg.set_ylabel('Ay [g]')


            ax_lowg.set_xlim(0, max(df["Steer_SW"].max() * 1.1, 0.6))

            ax_lowg.set_ylim(0, max(df["Ay"].max() * 1.1, y_plot.max() + 0.5))

            make_legend_outside()


            return round(k, DP), fig_lowg


        def calc_vda():

            result_container = {

                "v_linear": 0,

                "fig": None

            }


            def target_task():

                temp_filename = ""

                try:

                    temp_filename = os.path.abspath(f"temp_vda_{car_name}_{datetime.now().strftime('%H%M%S%f')}.csv")

                    target_cols = ["Time", "Steer_SW", "Ay", "M_SW", "Vx"]

                    existing_cols = [col for col in target_cols if col in df.columns]


                    df_temp = df[existing_cols].copy()

                    if 'Ay' in df_temp.columns:

                        df_temp['Ay'] = df_temp['Ay'] * 9.81

                    df_temp.to_csv(temp_filename, index=False)

                    cols = list(df_temp.columns)


                    xcxjsd_channel_list = [

                        safe_idx("TimeStep", cols),

                        safe_idx("Vx", cols),

                        safe_idx("Steer_SW", cols),

                        safe_idx("Ay", cols),

                        safe_idx("M_SW", cols),


                    ]

                    # vda

                    res_dict = xcxjsd_proc([temp_filename], xcxjsd_channel_list)

                    if res_dict and '平均值' in res_dict and len(res_dict['平均值']) > 0:


                        avg_data = res_dict['平均值']

                        result_container['v_linear'] = round(avg_data[0].get('1-线性区转向灵敏度(g/100deg)', 0), DP)

                    else:

                        print(f"[VDA] {car_name}计算返回结果为空")


                except Exception as e:

                    print(f"[VDA] 线程内计算出错: {e}")

                finally:

                    if temp_filename and os.path.exists(temp_filename):

                        try:

                            os.remove(temp_filename)

                        except:

                            pass


            # 线程

            t = threading.Thread(target=target_task)

            t.daemon = True

            t.start()


            t.join(timeout=50)


            if t.is_alive():

                print(f"[VDA] 警告: 计算超时 (超过50s)，已跳过！")


            return tuple(result_container.values())


        vda_linear, vda_fig = calc_vda()

        # lb_max_avx, lb_ay_time, lb_avz_over, lb_max_avx_fig, lb_ay_fig, lb_over_fig = calc_lb()


        my_k, my_fig = calc_my()

        metrics['线性区转向灵敏度'] = [vda_linear, my_k]

        figs['线性区转向灵敏度'] = [my_fig]


        return {'metrics': metrics, 'figs': figs}


    def _calc_step(self, df, car_name):

        metrics = {}

        figs = {}


        def calc_my():

            if 'Avx' in df.columns and 'TimeStep' in df.columns:


                my_max_avx = df["Avx"].max()


                return round(my_max_avx, DP)


        def calc_vda():

            result_container = {"max_avx": 0, "ay_time": 0, "avz_over": 0}


            def target_task():

                temp_filename = ""

                try:

                    temp_filename = os.path.abspath(f"temp_vda_{car_name}_{datetime.now().strftime('%H%M%S%f')}.csv")

                    target_cols = ["TimeStep", "Vx", "Steer_SW", "Ay", "Avz", "Roll"]

                    existing_cols = [col for col in target_cols if col in df.columns]


                    df_temp = df[existing_cols].copy()

                    if 'Ay' in df_temp.columns:

                        df_temp['Ay'] = df_temp['Ay'] * 9.8

                    df_temp.to_csv(temp_filename, index=False)

                    cols = list(df_temp.columns)


                    jy_channel_list = [

                        safe_idx("TimeStep", cols),

                        safe_idx("Vx", cols),

                        safe_idx("Steer_SW", cols),

                        safe_idx("Ay", cols),

                        safe_idx("Avz", cols),

                        safe_idx("Roll", cols),

                    ]


                    # vda

                    res_dict = jy_proc([temp_filename], jy_channel_list)

                    avg_data = res_dict['平均值']

                    vda_max_avx = round(avg_data[0].get('18-最大侧倾角速度(deg/s)', 0), DP)

                    vda_ay_time = round(avg_data[0].get('14-侧向加速度响应时间(s)', 0), DP)

                    vda_avz_over = round(avg_data[0].get('15-横摆角速度超调量(%)', 0), DP)


                    result_container["max_avx"] = vda_max_avx

                    result_container["ay_time"] = vda_ay_time

                    result_container["avz_over"] = vda_avz_over


                except Exception as e:

                    print(f"[VDA] 线程内计算出错: {e}")

                finally:

                    if temp_filename and os.path.exists(temp_filename):

                        try:

                            os.remove(temp_filename)

                        except:

                            pass


            # 线程

            t = threading.Thread(target=target_task)

            t.daemon = True

            t.start()


            t.join(timeout=50)


            # 5. 检查是否超时

            if t.is_alive():

                print(f"[VDA] 警告: 计算超时 (超过50s)，已跳过！")


            return tuple(result_container.values())

        def calc_lb():

            lb_max_avx = 0

            lb_ay_time = 0

            lb_avz_over = 0


            # 初始化 intercept 和 slope 默认值，防止绘图时报错

            slope = 0

            intercept = 0

            fit_success = False  # 标记是否拟合成功


            window_size = 100

            t_start = 10.05


            # --- 1. 计算侧向加速度响应时间 ---

            ay_segment = df['Ay'].iloc[1500:2501]  # 稳态ay

            ay_stable = ay_segment.mean()  # 稳态值


            ay_stable_90 = 0.9 * ay_stable

            # print(f"lbay90 {ay_stable_90}")

            rise_mask = (df['Ay'] >= 0.85 * ay_stable) & (df['Ay'] <= 0.95 * ay_stable)


            rise_data_ay = df.loc[rise_mask, 'Ay']

            rise_data_time = df.loc[rise_mask, 'TimeStep']


            if len(rise_data_ay) > 1:

                sorted_idx = np.argsort(rise_data_ay.values)

                t_end = np.interp(

                    ay_stable_90,

                    rise_data_ay.values[sorted_idx],

                    rise_data_time.values[sorted_idx]

                )

                lb_ay_time = round(t_end - t_start, DP)

                # print(f"lb_t1{t_start}")

                # print(f"lb_t2{t_end}")

            else:

                print(f"\t[{car_name}] LB侧加响应时间：无法计算 (数据不足)")

                t_end = t_start


            # --- 2. 计算侧倾角速度 (拟合) ---

            # 筛选 Roll 在 0.2~0.8 之间的数据

            roll_mask = (df['Roll'] < 0.8) & (df['Roll'] > 0.2)

            roll_x = df.loc[roll_mask, 'TimeStep']

            roll_y = df.loc[roll_mask, 'Roll']


            if len(roll_x) > 1:

                slope, intercept = np.polyfit(roll_x, roll_y, 1)

                lb_max_avx = round(slope, DP)

                fit_success = True  # 标记成功

            else:

                print(f"\t[{car_name}] LB侧倾角速度：无法计算 (数据不足)")


            # --- 3. 计算横摆角速度超调 ---

            y_avz_std = df['Avz'].rolling(window=window_size, center=True, min_periods=1).std()

            avz_condition = y_avz_std < (y_avz_std.max() * 0.1)


            if avz_condition.any():

                avz_stable_val = df.loc[avz_condition, 'Avz'].max()

            else:

                avz_stable_val = df['Avz'].iloc[-1]


            if avz_stable_val != 0:

                avz_max = df['Avz'].max()

                avz_max_idx = df['Avz'].idxmax()

                avz_max_time = df.loc[avz_max_idx, 'TimeStep']

                lb_avz_over = round((avz_max - avz_stable_val) / avz_stable_val * 100, DP)

            else:

                avz_max = 0

                avz_max_time = 0

                print(f"\t[{car_name}] LB横摆角速度超调量：无法计算 (稳态值为0)")


            # =========================================

            # 绘图 1: Roll 拟合图

            # =========================================

            fig_roll, ax = plt.subplots()

            # 画原始数据

            ax.plot(df['TimeStep'], df['Roll'], color='gray', alpha=0.5, linewidth=1, label='Raw Data (Full)')


            if fit_success:

                ax.scatter(roll_x, roll_y, color='blue', s=5, label='Fit Range Data', zorder=5)

                y_fitted_line = slope * roll_x + intercept


                ax.plot(roll_x, y_fitted_line, color='red', linewidth=2, linestyle='--',

                        label=f'Fit Line (k={slope:.2f})', zorder=10)


            ax.set_title(f'[方法2]-侧倾角速度-{car_name}')

            ax.set_xlabel('Time (s)')

            ax.set_ylabel('Roll (deg)')

            make_legend_outside()


            fig_ay, ax = plt.subplots()

            # 画原始数据

            ax.plot(df['TimeStep'], df['Ay'], color='gray', alpha=0.5, linewidth=1, label='Raw Data (Full)')


            if lb_ay_time != 0:

                y_start = np.interp(t_start, df['TimeStep'], df['Ay'])

                y_end = np.interp(t_end, df['TimeStep'], df['Ay'])

                ax.scatter([t_start, t_end], [y_start, y_end], color='orange', s=20, zorder=15)

                ax.axhline(y=ay_stable_90, color='black', linestyle=':', linewidth=1.5,

                           label=f'Steady State: {ay_stable_90:.2f} g')

                ax.annotate(f"{t_start:.3f}s",  # 显示的文本 (保留3位小数)

                            xy=(t_start, y_start),  # 箭头指向的点 (数据的真实坐标)

                            xytext=(10, -10),  # 文字的位置 (相对偏移量：向右10点，向下10点)

                            textcoords='offset points',  # 声明 xytext 使用的是相对偏移

                            ha='left',  # 水平对齐：左对齐

                            fontsize=8,  # 字号

                            color='darkorange')  # 字体颜色


                # 标注 t2

                ax.annotate(f"{t_end:.3f}s",

                            xy=(t_end, y_end),

                            xytext=(10, -10),  # 向右下方偏移，避免遮挡曲线

                            textcoords='offset points',

                            ha='left',

                            fontsize=8,

                            color='darkorange')


                y_range = df['Ay'].max() - df['Ay'].min()

                y_text_pos = df['Ay'].max() - y_range * 0.1

                ax.text((t_end + t_start) / 2, y_text_pos, f"dt={lb_ay_time}s", ha='center', va='bottom',

                        color='black', fontsize=10, fontweight='bold')


            ax.set_title(f'[方法2]-侧加响应时间-{car_name}')

            ax.set_xlabel('Time (s)')

            ax.set_ylabel('Ay (g)')

            make_legend_outside()


            # =========================================

            # 绘图 2: Avz 超调图

            # =========================================

            fig_avx, ax = plt.subplots()

            ax.plot(df['TimeStep'], df['Avz'], color='gray', alpha=0.5, linewidth=1, label='Raw Data (Full)')

            ax.axhline(y=avz_stable_val, color='black', linestyle=':', linewidth=1.5,

                       label=f'Steady State: {avz_stable_val:.2f} deg/s')


            if avz_stable_val != 0:

                ax.scatter(avz_max_time, avz_max, color='blue', s=20, label='Max Peak', zorder=5)  # 加大一点点点

                # 可以加一个箭头标注超调

                ax.annotate(f'Over: {lb_avz_over}%',

                            xy=(avz_max_time, avz_max),

                            xytext=(avz_max_time + 0.5, avz_max),

                            arrowprops=dict(facecolor='black', arrowstyle='->'))


            ax.set_title(f'[方法2]-超调量')

            ax.set_xlabel('Time (s)')

            ax.set_ylabel('Avz (deg/s)')

            make_legend_outside()


            return lb_max_avx, lb_ay_time, lb_avz_over, fig_roll, fig_ay, fig_avx



        # my_max_avx, my_ay_time, my_avz_over, my_max_avx_fig, my_ay_fig, my_over_fig = calc_my()

        # vda_max_avx, vda_ay_time, vda_avz_over = calc_vda()

        # lb_max_avx, lb_ay_time, lb_avz_over, lb_max_avx_fig, lb_ay_fig, lb_over_fig = calc_lb()

        #

        # metrics['侧倾角速度'] = [my_max_avx, vda_max_avx, lb_max_avx]

        # metrics['侧加响应时间'] = [my_ay_time, vda_ay_time, lb_ay_time]

        # metrics['横摆角速度超调量'] = [my_avz_over, vda_avz_over, lb_avz_over]

        #

        # figs['侧倾角速度'] = [my_max_avx_fig, lb_max_avx_fig]

        # figs['侧加响应时间'] = [my_ay_fig, lb_ay_fig]

        # figs['横摆角速度超调量'] = [my_over_fig, lb_over_fig]


        # my_max_avx = calc_my()

        vda_max_avx, vda_ay_time, vda_avz_over = calc_vda()

        lb_max_avx, lb_ay_time, lb_avz_over, lb_max_avx_fig, lb_ay_fig, lb_over_fig = calc_lb()


        metrics['侧倾角速度(deg/s)'] = [vda_max_avx, lb_max_avx]

        metrics['0.3g侧加响应时间(s)'] = [vda_ay_time, lb_ay_time]

        metrics['横摆角速度超调量(%)'] = [vda_avz_over, lb_avz_over]


        figs['侧倾角速度(deg/s)'] = [lb_max_avx_fig]

        figs['侧加响应时间(s)'] = [lb_ay_fig]

        figs['横摆角速度超调量(%)'] = [lb_over_fig]


        return {'metrics': metrics, 'figs': figs}


    def _calc_pulse(self, df, car_name):

        metrics = {}

        figs = {}


        def calc_vda():

            result_container = {"times": 0}


            def target_task():

                temp_filename = ""

                try:

                    temp_filename = os.path.abspath(f"temp_vda_{car_name}_{datetime.now().strftime('%H%M%S%f')}.csv")

                    target_cols = ["TimeStep", "Vx", "Ay", "Roll", "Avx"]

                    existing_cols = [col for col in target_cols if col in df.columns]


                    df_temp = df[existing_cols].copy()

                    if 'Ay' in df_temp.columns:

                        df_temp['Ay'] = df_temp['Ay'] * 9.8

                    df_temp.to_csv(temp_filename, index=False)

                    cols = list(df_temp.columns)


                    jy_channel_list = [

                        safe_idx("TimeStep", cols),

                        safe_idx("Vx", cols),

                        safe_idx("Ay", cols),

                        safe_idx("Roll", cols),

                        safe_idx("Avx", cols),

                    ]


                    # vda

                    res_dict = jmc_proc([temp_filename], jy_channel_list)

                    avg_data = res_dict['平均值']

                    vda_times = round(avg_data[0].get('12-晃动次数(次)', 0), DP)


                    result_container["times"] = vda_times


                except Exception as e:

                    print(f"[VDA] 线程内计算出错: {e}")

                finally:

                    if temp_filename and os.path.exists(temp_filename):

                        try:

                            os.remove(temp_filename)

                        except:

                            pass


            # 线程

            t = threading.Thread(target=target_task)

            t.daemon = True

            t.start()


            t.join(timeout=50)


            # 5. 检查是否超时

            if t.is_alive():

                print(f"[VDA] 警告: 计算超时 (超过50s)，已跳过！")


            return result_container["times"]


        vda_times = calc_vda()


        metrics['120km/h侧倾晃动次数(次)'] = [vda_times]


        # figs['120km/h侧倾晃动次数(次)'] = [lb_max_avx_fig]


        return {'metrics': metrics, 'figs': figs}

    def _calc_brake(self, df, car_name):

        """

        制动分析 - 双轴聚集点提取法 (X, Y均取实际数据的众数)

        彻底解决指令与实际值偏移的问题

        """

        metrics = {}

        figs = {}

        DP = 3


        # 1. 检查列

        if not all(col in df.columns for col in ['TimeStep', 'Pitch', 'Ax', 'Ax_SCcmd']):

            return {'metrics': metrics, 'figs': figs}


        # ==========================================

        # 2. 初始零点提取 (坐标点：[Ax_0, Pitch_0])

        # ==========================================

        base_mask = (df['Ax_SCcmd'].abs() < 0.01) & (df['TimeStep'] < 5.0)

        if base_mask.any():

            ax_0 = df['Ax'][base_mask].mean()

            pitch_0 = df['Pitch'][base_mask].mean()

        else:

            ax_0, pitch_0 = df['Ax'].iloc[0], df['Pitch'].iloc[0]


        # ==========================================

        # 3. 核心算法：双轴直方图寻优

        # ==========================================

        df['step_id'] = (df['Ax_SCcmd'].diff().abs() > 0.01).cumsum()


        x_points = [abs(ax_0)]

        y_points = [pitch_0]


        for _, group in df.groupby('step_id'):

            if abs(group['Ax_SCcmd'].mean()) < 0.05: continue


            # 提取当前段数据

            ax_data = group['Ax'].abs().values

            pitch_data = group['Pitch'].values

            if len(ax_data) < 20: continue


            # --- 寻找 Ax 的聚集点 (X轴) ---

            ax_bins = np.arange(ax_data.min(), ax_data.max() + 0.002, 0.002)

            if len(ax_bins) >= 2:

                counts_x, edges_x = np.histogram(ax_data, bins=ax_bins)

                ax_ss = (edges_x[np.argmax(counts_x)] + edges_x[np.argmax(counts_x) + 1]) / 2

            else:

                ax_ss = ax_data.mean()


            # --- 寻找 Pitch 的聚集点 (Y轴) ---

            p_bins = np.arange(pitch_data.min(), pitch_data.max() + 0.005, 0.005)

            if len(p_bins) >= 2:

                counts_y, edges_y = np.histogram(pitch_data, bins=p_bins)

                pitch_ss = (edges_y[np.argmax(counts_y)] + edges_y[np.argmax(counts_y) + 1]) / 2

            else:

                pitch_ss = pitch_data.mean()


            x_points.append(ax_ss)

            y_points.append(pitch_ss)


        # ==========================================

        # 4. 计算梯度与绘图

        # ==========================================

        slope, intercept, r_value, _, _ = linregress(x_points, y_points)


        fig, ax_plot = plt.subplots()


        # 绘背景轨迹

        plot_mask = df['Vx'] > 1.0 if 'Vx' in df.columns else slice(None)

        ax_plot.scatter(df['Ax'][plot_mask].abs(), df['Pitch'][plot_mask],

                        c='gray', s=5, alpha=0.1, label='Actual Trajectory')


        ax_plot.scatter(x_points, y_points, color='red', s=80, edgecolors='black',

                        zorder=10, label='稳态聚集点')


        # 绘拟合线

        line_x = np.array([0, max(x_points) * 1.1])

        ax_plot.plot(line_x, slope * line_x + intercept, 'r--',

                     label=f'Fit: {slope:.3f} deg/g)')


        ax_plot.set_title(f'[稳态聚集点拟合]{car_name} Pitch Gradient')

        ax_plot.set_xlabel('|Ax| (g)')

        ax_plot.set_ylabel('Pitch (deg)')

        ax_plot.grid(True, linestyle=':', alpha=0.6)


        make_legend_outside()


        metrics['制动总俯仰梯度(-)'] = [round(slope, DP)]

        figs['制动总俯仰梯度(-)'] = fig


        return {'metrics': metrics, 'figs': figs}


    def _calc_5kmh_brake(self, df, car_name):

        metrics = {}

        figs = {}


        def calc_vda():

            result_container = {

                "v_chuandong": 0,

                "v_diantoujiao": 0,

                "v_yangjiao": 0,

                'v_fuyang': 0,

                "fig": None

            }


            def target_task():

                temp_filename = ""

                try:

                    temp_filename = os.path.abspath(f"temp_vda_{car_name}_{datetime.now().strftime('%H%M%S%f')}.csv")

                    target_cols = ["TimeStep", "Vx", "Ax", "Pitch"]

                    existing_cols = [col for col in target_cols if col in df.columns]


                    df_temp = df[existing_cols].copy()

                    if 'Ax' in df_temp.columns:

                        df_temp['Ax'] = df_temp['Ax'] * 9.8

                        df_temp['Pitch'] = df_temp['Pitch'] * -1

                    df_temp.to_csv(temp_filename, index=False)

                    cols = list(df_temp.columns)


                    qlzd_channel_list = [

                        safe_idx("TimeStep", cols),

                        safe_idx("Vx", cols),

                        safe_idx("Ax", cols),

                        safe_idx("Pitch", cols),

                    ]


                    res_dict = qlzd_proc([temp_filename], qlzd_channel_list)


                    if res_dict and '平均值' in res_dict and len(res_dict['平均值']) > 0:

                        avg_data = res_dict['平均值']


                        result_container['v_diantoujiao'] = round(avg_data[0].get('11-最大俯角(deg)', 0), DP)

                        result_container['v_yangjiao'] = round(avg_data[0].get('12-最大仰角(deg)', 0), DP)

                        result_container['v_fuyang'] = round(avg_data[0].get('13-俯仰收敛次数(次)', 0), DP)

                        result_container['v_chuandong'] = round(avg_data[0].get('14-制动窜动次数(次)', 0), DP)


                        # if len(avg_data) > 1 and isinstance(avg_data[1], dict):

                        #     plots_dict = avg_data[1]

                        #     target_key = '方向盘转角-侧向加速度'

                        #

                        #     if target_key in plots_dict:

                        #         plot_data = plots_dict[target_key]

                        #

                        #         fig_vda, ax_v = plt.subplots()

                        #         # 遍历 curve 列表画图

                        #         if 'curve' in plot_data:

                        #             for curve_info in plot_data['curve']:

                        #                 x_data = curve_info.get('x')

                        #                 y_data = curve_info.get('y')

                        #                 name = curve_info.get('curvename', 'Curve')

                        #                 ctype = curve_info.get('curvetype', 'line')

                        #

                        #                 # 确保数据有效

                        #                 if x_data is not None and y_data is not None and len(x_data) > 0:

                        #                     # 1. 如果是原始数据 -> 灰色半透明散点，作为背景

                        #                     if '原始' in name:

                        #                         ax_v.scatter(x_data, y_data, s=3, alpha=0.2, color='gray', label=name,

                        #                                      zorder=1)

                        #

                        #                     # 2. 如果是拟合数据 -> 彩色加粗实线，区分左右转

                        #                     elif '拟合' in name:

                        #                         c = 'red' if '左' in name else ('blue' if '右' in name else 'green')

                        #                         ax_v.plot(x_data, y_data, lw=2.5, color=c, label=name, zorder=5)

                        #

                        #                     # 3. 兜底或其他类型

                        #                     else:

                        #                         if ctype == 'scatter':

                        #                             ax_v.scatter(x_data, y_data, s=4, alpha=0.5, label=name)

                        #                         else:

                        #                             ax_v.plot(x_data, y_data, lw=1.5, linestyle='--', label=name)

                        #

                        #         ax_v.set_title(f"[VDA] Ay vs Steer")

                        #         ax_v.set_xlabel(plot_data.get('xlabel', 'Steer [deg]'))

                        #         ax_v.set_ylabel(plot_data.get('ylabel', 'Ay [m/s^2]'))

                        #         make_legend_outside()

                        #         result_container['fig'] = fig_vda

                    else:

                        print(f"[VDA] {car_name}计算返回结果为空")


                except Exception as e:

                    print(f"[VDA] 线程内计算出错: {e}")

                finally:

                    if temp_filename and os.path.exists(temp_filename):

                        try:

                            os.remove(temp_filename)

                        except:

                            pass


            # 线程

            t = threading.Thread(target=target_task)

            t.daemon = True

            t.start()


            t.join(timeout=50)


            if t.is_alive():

                print(f"[VDA] 警告: 计算超时 (超过50s)，已跳过！")


            return tuple(result_container.values())


        def calc_my():

            t_start_mask = df['TimeStep'] < 4.0

            pitch_zero_ref = df.loc[t_start_mask, 'Pitch'].mean()

            # 去掉零漂移

            df['Pitch_Relative'] = df['Pitch'] - pitch_zero_ref

            my_diantoujiao = df['Pitch_Relative'].max()

            my_yangjiao= df['Pitch_Relative'].min() *-1


            return round(my_diantoujiao,DP), round(my_yangjiao, DP)


        v_cuandong, v_diantoujiao, v_yangjiao, v_fuyang, fig = calc_vda()

        my_diantoujiao, my_yangjiao = calc_my()


        metrics['全力制动窜动次数(次)'] = [v_cuandong]

        metrics['全力制动最大点头角(俯角,deg)'] = [v_diantoujiao, my_diantoujiao]

        metrics['全力制动最大仰角(deg)'] = [v_yangjiao, my_yangjiao]

        metrics['全力制动俯仰次数(次)'] = [v_fuyang]


        return {'metrics': metrics, 'figs': figs}


    def _calc_acc(self, df, car_name):

        metrics = {}

        figs = {}


        def calc_my():

            dt = np.mean(np.diff(df["TimeStep"]))

            fs = 1.0 / dt if dt > 0 else 100.0


            mask_zero = (df["TimeStep"] >= 0) & (df["TimeStep"] <= 5)

            if mask_zero.sum() > 0:

                pitch_zero_drift = df.loc[mask_zero, "Pitch"].mean()

            else:

                pitch_zero_drift = 0.0  # 容错：如果没有0-5s的数据则不扣除


            pitch_corrected = df["Pitch"] - pitch_zero_drift


            # 第一步：中值滤波 (Median Filter)

            kernel_size = int(fs * 0.8)

            if kernel_size % 2 == 0:

                kernel_size += 1

            pitch_despiked = medfilt(pitch_corrected, kernel_size=kernel_size)


            # 第二步：低通滤波 (Low-pass) - 抚平微小褶皱

            cutoff = 5

            nyq = 0.5 * fs

            normal_cutoff = cutoff / nyq

            b, a = butter(4, normal_cutoff, btype='low', analog=False)

            pitch_filtered = filtfilt(b, a, pitch_despiked)

            # --------------------------------------------------


            # --- 2. 数据截取 ---

            # 定义 5 个加速段的稳态时间窗口

            time_windows = [

                (6, 24),  # 0.05g段

                (56, 74),  # 0.10g段

                (106, 124),  # 0.15g段

                (156, 174),  # 0.20g段

                (206, 224)  # 0.25g段

            ]


            nominal_ax_targets = [0.05, 0.10, 0.15, 0.20, 0.25]


            ax_steady = []

            pitch_steady = []


            # 遍历每个时间窗口和对应的标称加速度

            for (t_start, t_end), target_ax in zip(time_windows, nominal_ax_targets):

                mask = (df["TimeStep"] >= t_start) & (df["TimeStep"] <= t_end)


                if mask.sum() > 0:

                    # 不再取实际的 Ax_mean，直接使用 target_ax

                    pitch_mean = np.mean(pitch_filtered[mask]) * -1


                    ax_steady.append(target_ax)

                    pitch_steady.append(pitch_mean)


            if len(ax_steady) < 2:

                print(f"\t[稳态聚集点拟合] 加速俯仰梯度：有效稳态数据点不足，无法拟合")

                return 0, None


            x_fit = np.array(ax_steady)

            y_fit = np.array(pitch_steady)


            # 线性拟合 (一次多项式)

            p_coeffs = np.polyfit(x_fit, y_fit, 1)

            k = p_coeffs[0]  # 梯度：deg/g


            # --- 3. 绘图部分 (1行2列的并排子图) ---

            fig, (ax_time, ax_pitch) = plt.subplots(2, 1)


            pitch_raw_disp = pitch_corrected * -1

            pitch_filt_disp = pitch_filtered * -1


            # 【左侧子图：时域滤波对比与截取窗口】

            ax_time.plot(df["TimeStep"], pitch_raw_disp, color='gray', alpha=0.4, label='Raw Pitch')

            ax_time.plot(df["TimeStep"], pitch_filt_disp, color='red', linewidth=1.5, label='Filtered Pitch')


            # 画出绿色透明矩形框，明确展示数据是在哪些时间段内提取的

            for i, (t_start, t_end) in enumerate(time_windows):

                label = 'Steady Windows' if i == 0 else ""

                ax_time.axvspan(t_start, t_end, color='green', alpha=0.15, label=label)


            ax_time.set_title(f'[稳态聚集点拟合] {car_name}Pitch 滤波&提取窗')

            ax_time.set_xlabel('Time [s]')

            ax_time.set_ylabel('Pitch * -1 [deg]')


            ax_pitch.scatter(df["Ax"], pitch_raw_disp, color='gray', s=3, alpha=0.15, label='Raw Data (All)')

            ax_pitch.scatter(x_fit, y_fit, color='black', s=45, marker='o', alpha=0.9,

                             label='Steady-state Points', zorder=5)


            # 绘制拟合线

            x_plot = np.linspace(0, max(x_fit) * 1.1, 50)

            y_plot = np.polyval(p_coeffs, x_plot)

            ax_pitch.plot(x_plot, y_plot, 'r-', linewidth=2.5, label=f'Linear Fit (K={k:.2f} deg/g)', zorder=4)


            ax_pitch.set_title(f'[稳态聚集点拟合] {car_name} 加速稳态俯仰梯度\n K = {k:.3f} deg/g')

            ax_pitch.set_xlabel('目标Ax[g]')

            ax_pitch.set_ylabel('Pitch[deg]')


            ax_pitch.set_xlim(0, max(max(x_fit) * 1.2, 0.3))

            y_min = min(pitch_raw_disp.min(), min(y_plot) - 0.5)

            y_max = max(pitch_raw_disp.max(), max(y_plot) + 0.5)

            ax_pitch.set_ylim(y_min, y_max)


            # 自动调整子图间距防重叠

            plt.tight_layout()


            make_legend_outside()


            return round(k, DP), fig


        my_k, my_fig = calc_my()

        metrics['加速俯仰总梯度(-)'] = [my_k]

        figs['加速俯仰总梯度(-)'] = [my_fig]


        return {'metrics': metrics, 'figs': figs}

    # def _calc_acc(self, df, car_name):

    #     """

    #     加速分析 - 自动识别加速阶梯工况并计算俯仰梯度 (deg/g)

    #     支持从 0.05g 开始的自动扫掠提取

    #     """

    #     metrics = {}

    #     figs = {}

    #

    #     # 1. 检查必需的列

    #     required_cols = ['TimeStep', 'Pitch', 'Ax', 'Ax_SCcmd']

    #     if not all(col in df.columns for col in required_cols):

    #         print(f"⚠️ [{car_name}] 缺少加速分析所需列 {required_cols}，已跳过。")

    #         return {'metrics': metrics, 'figs': figs}

    #

    #     time = df['TimeStep']

    #     pitch = df['Pitch']

    #     ax = df['Ax']

    #     ax_cmd = df['Ax_SCcmd']

    #

    #     # ==========================================

    #     # 2. 初始零点提取 (取 Ax_SCcmd == 0 且静止/低速前期)

    #     # ==========================================

    #     base_mask = (ax_cmd.abs() < 0.01) & (time < 2.0)

    #     pitch_0 = pitch[base_mask].mean() if base_mask.any() else pitch.iloc[0]

    #     ax_0 = ax[base_mask].mean() if base_mask.any() else ax.iloc[0]

    #

    #     # ==========================================

    #     # 3. 自动识别加速阶梯

    #     # ==========================================

    #     # 识别指令变化

    #     df['step_id'] = (ax_cmd.diff().abs() > 0.01).cumsum()

    #

    #     step_results = []

    #

    #     for _, group in df.groupby('step_id'):

    #         cmd_val = group['Ax_SCcmd'].mean()

    #

    #         # 只处理真正的加速阶梯 (Ax_SCcmd > 0.05g)

    #         if cmd_val < 0.05:

    #             continue

    #

    #         # 【关键：稳态选取逻辑】

    #         # 加速时后期可能动力不足导致掉g，且前期有换挡冲击

    #         # 建议取该段数据的 40% - 80% 区间，避开末尾的动力衰减

    #         n = len(group)

    #         if n < 20: continue  # 过滤掉过短的过渡段

    #

    #         steady_part = group.iloc[int(n * 0.4): int(n * 0.8)]

    #

    #         # 动力充足检查：如果实际 Ax 与指令差距过大（如超过 0.05g），说明动力饱和，该点舍弃

    #         actual_ax = steady_part['Ax'].mean()

    #         if abs(actual_ax - cmd_val) > 0.05:

    #             # print(f"ℹ️ [{car_name}] 指令 {cmd_val:.2f}g 动力不足(实际{actual_ax:.2f}g)，已舍弃该点")

    #             continue

    #

    #         if len(steady_part) > 5:

    #             step_results.append({

    #                 'ax_ss': actual_ax,

    #                 'pitch_ss': steady_part['Pitch'].mean()

    #             })

    #

    #     if len(step_results) < 1:

    #         print(f"⚠️ [{car_name}] 有效加速阶梯不足，无法计算梯度。")

    #         return {'metrics': metrics, 'figs': figs}

    #

    #     # ==========================================

    #     # 4. 计算梯度 (线性回归)

    #     # ==========================================

    #     # 包含零点

    #     x_points = [abs(ax_0)] + [abs(res['ax_ss']) for res in step_results]

    #     y_points = [pitch_0] + [res['pitch_ss'] for res in step_results]

    #

    #     slope, intercept, r_value, p_value, std_err = linregress(x_points, y_points)

    #     pitch_gradient = slope

    #

    #     # ==========================================

    #     # 5. 绘图可视化

    #     # ==========================================

    #     fig_accel, ax_plot = plt.subplots(figsize=(8, 6))

    #

    #     # 绘制全量数据 (过滤掉车速过高或回位段)

    #     plot_mask = (df['Vx'] < 120) & (df['Ax'] > -0.05) if 'Vx' in df.columns else slice(None)

    #     ax_plot.scatter(ax[plot_mask].abs(), pitch[plot_mask], c='gray', s=2, alpha=0.15, label='Transient Data')

    #

    #     # 绘制稳态点

    #     ax_plot.scatter(x_points, y_points, color='blue', s=60, zorder=5, label='Steady State Points')

    #

    #     # 绘制拟合线

    #     line_x = np.array([0, max(x_points) * 1.2])

    #     line_y = slope * line_x + intercept

    #     ax_plot.plot(line_x, line_y, 'b--', linewidth=2, label=f'Fit: {slope:.3f} deg/g (R²={r_value ** 2:.3f})')

    #

    #     ax_plot.set_title(f'[AOS]{car_name}Acceleration Pitch Gradient')

    #     ax_plot.set_xlabel('Acceleration |Ax| (g)')

    #     ax_plot.set_ylabel('Pitch (deg)')

    #     ax_plot.grid(True, linestyle=':', alpha=0.6)

    #

    #     # 这里调用你外部定义的 legend 函数

    #     if 'make_legend_outside' in globals():

    #         make_legend_outside(ax_plot)

    #     else:

    #         ax_plot.legend()

    #

    #     # 填充结果

    #     metrics['加速总俯仰梯度'] = [round(pitch_gradient, DP)]

    #

    #     figs['加速总俯仰梯度'] = fig_accel

    #

    #     return {'metrics': metrics, 'figs': figs}

    def _clac_bump(self, df, car_name):

        metrics = {}

        figs = {}


        def calc_vda():

            result_container = {

                "v_L1": 0,

                "v_L2": 0,

                "v_avy": 0,

                'v_az': 0,

                "fig": None

            }


            def target_task():

                temp_filename = ""

                try:

                    temp_filename = os.path.abspath(f"temp_vda_{car_name}_{datetime.now().strftime('%H%M%S%f')}.csv")

                    target_cols = ["TimeStep", "Vx", "Jnc_L1", "Jnc_L2", "Avy", "Az"]

                    existing_cols = [col for col in target_cols if col in df.columns]


                    df_temp = df[existing_cols].copy()

                    if 'Az' in df_temp.columns:

                        df_temp['Az'] = df_temp['Az'] * 9.8

                    df_temp.to_csv(temp_filename, index=False)

                    cols = list(df_temp.columns)


                    qifu_channel_list = [

                        safe_idx("TimeStep", cols),

                        safe_idx("Vx", cols),

                        safe_idx("Jnc_L1", cols),

                        safe_idx("Jnc_L2", cols),

                        safe_idx("Avy", cols),

                        safe_idx("Az", cols),

                    ]


                    res_dict = qifu_proc([temp_filename], qifu_channel_list)


                    if res_dict and '平均值' in res_dict and len(res_dict['平均值']) > 0:

                        avg_data = res_dict['平均值']


                        result_container['v_L1'] = round(avg_data[0].get('11-前轴最大位移(mm)', 0), DP)

                        result_container['v_L2'] = round(avg_data[0].get('12-后轴最大位移(mm)', 0), DP)

                        result_container['v_avy'] = round(avg_data[0].get('15-最大俯仰角速度(deg/s)', 0), DP)

                        result_container['v_az'] = round(avg_data[0].get('16-最大垂向加速度(m/s2)', 0)/9.81, DP)


                    else:

                        print(f"[VDA] {car_name}计算返回结果为空")


                except Exception as e:

                    print(f"[VDA] 线程内计算出错: {e}")

                finally:

                    if temp_filename and os.path.exists(temp_filename):

                        try:

                            os.remove(temp_filename)

                        except:

                            pass


            # 线程

            t = threading.Thread(target=target_task)

            t.daemon = True

            t.start()


            t.join(timeout=50)


            if t.is_alive():

                print(f"[VDA] 警告: 计算超时 (超过50s)，已跳过！")


            return tuple(result_container.values())



        v_L1, v_L2, v_AVY, v_AZ, fig = calc_vda()


        metrics['前轴垂向位移(mm)'] = [v_L1]

        metrics['后轴垂向位移(mm)'] = [v_L2]

        metrics['俯仰角速度(deg/s)'] = [v_AVY]

        metrics['车身垂向加速度(g)'] = [v_AZ]


        return {'metrics': metrics, 'figs': figs}


    def _calc_generic(self, df, cond_cfg, car_name):

        """通用计算逻辑"""

        metrics = {}

        for ind in cond_cfg.get('indicators', []):

            label = ind['label']

            y_cfg = ind.get('y1', {})

            col = y_cfg.get('name') if isinstance(y_cfg, dict) else y_cfg

            if col in df.columns:

                op = ind.get('op')

                data = df[col]

                if isinstance(y_cfg, dict):

                    if y_cfg.get('scale'): data = data * y_cfg['scale']

                    if y_cfg.get('sub_ori'): data = data - data.iloc[0]


                if op == 'y1_max':

                    val = data.max()

                elif op == 'y1_abs_max':

                    val = data.abs().max()

                else:

                    val = "-"

                metrics[label] = round(val, 4) if isinstance(val, (int, float)) else val

        return {'metrics': metrics, 'figs': {}}


    def _calc_multi_condition(self, car_name, car_folder_path, common_cfg):

        """专项：跨多工况的 ZMP (零力矩点) 侧翻风险综合评估"""

        metrics = {}

        figs = {}


        if not car_folder_path or not os.path.exists(car_folder_path):

            return {'metrics': metrics, 'figs': figs}


        # 1. 寻找需要的多个工况 CSV 文件 (比如阶跃和稳态，你可以根据实际需要改名字)

        path_steady = find_file_fuzzy(car_folder_path, "稳态回转")

        path_central = find_file_fuzzy(car_folder_path, "中心区")


        df_steady_raw = safe_read_csv(path_steady) if path_steady else None

        df_central_raw = safe_read_csv(path_central) if path_central else None


        df_steady = self._preprocess_df(df_steady_raw, common_cfg) if df_steady_raw is not None else None

        df_central = self._preprocess_df(df_central_raw, common_cfg) if df_central_raw is not None else None


        # 2. 获取车辆几何参数

        veh_params = self.veh_info_map.get(car_name, {})

        f_wheel_base = veh_params.get('f_wheel_base', 1.0)

        r_wheel_base = veh_params.get('r_wheel_base', 1.0)

        f_track_width = veh_params.get('f_track_width', 1.0)

        r_track_width = veh_params.get('r_track_width', 1.0)


        def calc_zmp_trajectory(df_target, cond_name):

            """计算单工况的 ZMP 轨迹并返回最大侧向偏移量"""

            if df_target is None or df_target.empty:

                return 0.0, None, None


            # 🚨这里必须确保你的 CSV 里有四轮垂向力通道！

            # 根据你 CarSim 输出的实际名称修改下面的字符串

            fz_cols = ['Fz_L1', 'Fz_R1', 'Fz_L2', 'Fz_R2']


            # 如果别名映射成了别的名字，比如 'Fz_fl', 请替换上面的列表

            if not all(col in df_target.columns for col in fz_cols):

                print(f"\t[{car_name}] {cond_name} 缺少垂向力通道 {fz_cols}，跳过 ZMP 计算")

                return 0.0, None, None


            Fz_fl = df_target[fz_cols[0]]

            Fz_fr = df_target[fz_cols[1]]

            Fz_rl = df_target[fz_cols[2]]

            Fz_rr = df_target[fz_cols[3]]


            # 总垂向力，防止四轮离地时除以 0 报错

            Fz_sum = Fz_fl + Fz_fr + Fz_rl + Fz_rr

            Fz_sum = Fz_sum.replace(0, 1e-6)


            # 计算局部坐标系下的 ZMP (原点为质心投影点)

            # X_zmp: 纵向零力矩点 (向车头为正)

            x_zmp = ((Fz_fl + Fz_fr) * f_wheel_base - (Fz_rl + Fz_rr) * r_wheel_base) / Fz_sum

            # Y_zmp: 横向零力矩点 (向左为正)

            y_zmp = ((Fz_fl + Fz_rl) * (f_track_width / 2) - (Fz_fr + Fz_rr) * (f_track_width / 2)) / Fz_sum


            max_y = y_zmp.abs().max()

            return max_y, x_zmp, y_zmp


        # 3. 分别计算各工况的 ZMP

        max_y_steady, x_zmp_steady, y_zmp_steady = calc_zmp_trajectory(df_steady, "Steady")

        max_y_step, x_zmp_step, y_zmp_step = calc_zmp_trajectory(df_central, "Central")


        # 找出最危险的工况 (侧倾最大，Y_zmp 最大)

        global_max_y = max(max_y_steady, max_y_step)


        # 换算成“侧翻危险指数 Rollover Index” (超过 100% 即代表一侧车轮完全离地)

        rollover_index = (global_max_y / (f_track_width / 2)) * 100


        metrics['最大横向ZMP偏移(m)'] = [round(global_max_y, DP)]

        metrics['侧翻危险指数(%)'] = [round(rollover_index, 1)]


        # 4. 可视化：画出 ZMP 散点与支撑多边形边界

        fig_zmp, ax = plt.subplots()


        # 画出车辆的“支撑多边形” (四个轮子接地点连成的矩形)

        rect_x = [-f_track_width / 2, f_track_width / 2, f_track_width / 2, -f_track_width / 2, -f_track_width / 2]

        rect_y = [f_wheel_base, f_wheel_base, -r_wheel_base, -r_wheel_base, f_wheel_base]

        ax.plot(rect_x, rect_y, 'k-', linewidth=2, label='支撑多边形边界 (Support Polygon)')

        ax.axhline(0, color='gray', linestyle='--', linewidth=0.5)

        ax.axvline(0, color='gray', linestyle='--', linewidth=0.5)


        # 把轨迹画上去 (只画有数据的)

        if y_zmp_steady is not None:

            ax.scatter(y_zmp_steady, x_zmp_steady, s=2, alpha=0.3, color='blue', label='Steady-State ZMP')

        if y_zmp_step is not None:

            ax.scatter(y_zmp_step, x_zmp_step, s=2, alpha=0.3, color='red', label='Step Steer ZMP')


        # 标注重心

        ax.plot(0, 0, 'go', markersize=8, label='质心 (CG)')


        ax.set_title(f'[综合防侧翻] {car_name} ZMP 轨迹\nMax Offset: {global_max_y:.3f}m | Risk: {rollover_index:.1f}%')

        ax.set_xlabel('横向偏移 Y (m)')

        ax.set_ylabel('纵向偏移 X (m)')

        # 强制比例一致，使得画出来的框是真实物理比例

        ax.axis('equal')

        ax.set_xlim(-f_track_width, f_track_width)

        ax.set_ylim(-r_wheel_base - 0.5, f_wheel_base + 0.5)


        # 如果越界，加上红色警告底色

        if rollover_index >= 100:

            ax.set_facecolor('#fff0f0')  # 极浅的红色

            ax.text(0, 0, "DANGER: 发生单侧车轮离地", color='red', fontsize=12, fontweight='bold', ha='center')


        try:

            make_legend_outside(ax)

        except:

            ax.legend()


        figs['最大横向ZMP偏移(m)'] = fig_zmp

        figs['侧翻危险指数(%)'] = fig_zmp


        return {'metrics': metrics, 'figs': figs}


    def get_veh_config_from_txt(self, txt_file_path):

        # 1. 初始化默认值，防止文件缺失或解析失败报错

        params = {

            "name": None,

            "category": None,

            "wheel_base": 0.0,

            "f_wheel_base": 0.0,

            "r_wheel_base": 0.0,

            "f_weight": 0.0,

            "r_weight": 0.0,

            "steer_ratio": 0.0

        }


        if not os.path.exists(txt_file_path):

            print(f"❌ 文件不存在: {txt_file_path}")

            return params


        try:

            with open(txt_file_path, 'r', encoding='utf-8') as f:

                for line in f:

                    line = line.strip()

                    # 跳过没有冒号的行（如分隔线、标题）

                    if ':' not in line:

                        continue


                    # 分割 Key 和 Value (只分割第一个冒号)

                    key_str, val_str = line.split(':', 1)

                    key = key_str.strip()

                    val = val_str.strip()

                    # --- 根据你写入的字符串进行精确匹配 ---


                    if key == "Dataset":

                        params["name"] = val


                    elif key == "Category":

                        params["category"] = val


                    elif key == "Wheel Base (m)":

                        params["wheel_base"] = float(val)


                    elif key == "Front Wheel Base (m)":

                        params["f_wheel_base"] = float(val)


                    elif key == "Rear Wheel Base (m)":

                        params["r_wheel_base"] = float(val)


                    elif key == "Front Weight (kg)":

                        params["f_weight"] = float(val)


                    elif key == "Rear Weight (kg)":

                        params["r_weight"] = float(val)


                    elif key == "Steer Ratio":

                        params["steer_ratio"] = float(val)


            print(f"✅ 参数读取成功: {params['name']} (SR={params['steer_ratio']})")

            print(f"✅ WHEELBAESE: {params['f_wheel_base']},{params['r_wheel_base']}")


        except Exception as e:

            print(f"❌ 读取 config.txt 发生错误: {e}")


        return params



# ==========================================

# 2. 报告详情页生成逻辑

# ==========================================


def create_comparison_plot(target_ind, cars_data_map, selected_cars):

    """

    [新增通用函数] 创建原始数据对比图

    支持: 单图叠加 (Default) 和 多图分列 (Split Mode)

    增加: 自动读取并显示单位

    """

    # 1. 解析配置 (列名 & 单位)

    x_cfg = target_ind.get('x', {})

    y_cfg = target_ind.get('y1', {})


    x_col = x_cfg.get('name', 'TimeStep')

    y_col = y_cfg.get('name')


    x_unit = x_cfg.get('unit', '')

    y_unit = y_cfg.get('unit', '')


    # 构造显示的标签文本，例如: "Roll [deg]"

    x_label = f"{x_col} [{x_unit}]" if x_unit else x_col

    y_label = f"{y_col} [{y_unit}]" if y_unit else y_col


    # 2. 检查是否有有效数据

    has_data = False

    for car in selected_cars:

        data_source = cars_data_map.get(car, {})

        df = data_source.get('data') if isinstance(data_source, dict) else data_source


        if df is not None and not df.empty and x_col in df.columns and y_col in df.columns:

            has_data = True

            break


    if not has_data:

        return None


    # 3. 检查显示模式

    split_mode = str(target_ind.get('split_plots', "False")) == "True"


    if split_mode:

        # --- 分图模式 ---

        num_cars = len(selected_cars)

        cols = 2

        rows = math.ceil(num_cars / cols)

        fig_h = max(4, rows * 3)

        fig, axes = plt.subplots(rows, cols, figsize=(6, fig_h))


        if rows * cols > 1:

            axes_flat = axes.flatten()

        else:

            axes_flat = [axes]


        for i, car in enumerate(selected_cars):

            ax = axes_flat[i]

            res = cars_data_map.get(car, {})

            df = res.get('data')


            if df is not None and not df.empty and x_col in df.columns and y_col in df.columns:

                ax.plot(df[x_col], df[y_col], linewidth=1.5, label=car, color='tab:blue')

                ax.set_title(car)

                # 使用带单位的标签

                ax.set_xlabel(x_label)

                ax.set_ylabel(y_label)

            else:

                ax.text(0.5, 0.5, "No Data", ha='center', va='center', transform=ax.transAxes)


        # 隐藏多余的子图

        for j in range(i + 1, len(axes_flat)):

            axes_flat[j].axis('off')


        fig.tight_layout()


    else:

        # --- 叠加模式 ---

        fig, ax = plt.subplots()


        for car in selected_cars:

            res = cars_data_map.get(car, {})

            df = res.get('data')

            if df is not None and not df.empty and x_col in df.columns and y_col in df.columns:

                ax.plot(df[x_col], df[y_col], label=car, linewidth=1.5)


        ax.set_title(target_ind.get('label', y_col))  # 如果没有label就用列名兜底

        # 使用带单位的标签

        ax.set_xlabel(x_label)

        ax.set_ylabel(y_label)


        # 图例处理 (只有叠加模式才需要图例)

        ax.legend()


    make_legend_outside()


    return fig



def create_detailed_pages(doc, car_folders, calc, configs, pre_calc_data=None):

    print("\n---------- 阶段 2: 生成详情页 (图表组合) ----------")


    # 定义边框样式

    border_heavy = {"sz": 18, "val": "single", "color": "000000"}

    border_thin = {"sz": 4, "val": "single", "color": "000000"}


    num_cars = len(car_folders)

    METHOD_WIDTH = 3.5  # 计算方案列宽

    car_col_width = (PAGE_USABLE_WIDTH - IND_WIDTH - METHOD_WIDTH) / num_cars if num_cars > 0 else 0

    col_widths_list = [IND_WIDTH, METHOD_WIDTH] + [car_col_width] * num_cars


    def set_table_layout(tbl, width_list):

        tbl.autofit = False

        tbl.allow_autofit = False


        tblPr = tbl._tblPr

        tblLayout = tblPr.first_child_found_in("w:tblLayout")

        if tblLayout is None:

            tblLayout = OxmlElement('w:tblLayout')

            tblPr.append(tblLayout)

        tblLayout.set(qn('w:type'), 'fixed')


        for col_idx, (col, width_cm) in enumerate(zip(tbl.columns, width_list)):

            col.width = Cm(width_cm)

            for cell in col.cells:

                cell.width = Cm(width_cm)


    header_font_size = 6 if len(car_folders) > 4 else 8


    for kw, cfg in configs.items():

        if kw == 'common_config': continue


        cars_res = {}

        for car in car_folders:

            # A. 优先从缓存取

            if pre_calc_data and kw in pre_calc_data and car in pre_calc_data[kw]:

                cars_res[car] = pre_calc_data[kw][car]

            else:

                cars_res[car] = calc.get_condition_results(car, kw, None, configs)


        has_valid_data = False

        for car, res in cars_res.items():

            # 检查 DataFrame 是否存在且不为空

            df = res.get('data')

            if df is not None and not df.empty:

                has_valid_data = True

                break


        # 如果所有车都没有数据，直接跳过这个工况，不生成标题也不生成图表

        if not has_valid_data:

            print(f"  [Info] 跳过详情页工况: {kw} (无有效数据)")

            continue


        # --- 工况分页与标题 ---

        doc.add_page_break()

        p = doc.add_paragraph()

        p.style = 'Heading 1'

        run = p.add_run(f"{cfg.get('title', kw)}")

        set_run_font(run, size_pt=14, bold=True)


        active_table = None


        for ind in cfg.get('indicators', []):

            label = ind['label']

            has_fig = str(ind.get('has_fig', "True")) == "True"

            # 3. 计算方案

            raw_method = ind.get('method', '-')

            method_list = [str(m) for m in raw_method] if isinstance(raw_method, list) else str(raw_method).split(

                '\n')


            # =======================================================

            # 分支 A: 如果有图 -> 开启新章节 (标题+说明+图+新表头)

            # =======================================================

            if has_fig:

                # 2. 指标标题

                p = doc.add_paragraph()

                p.paragraph_format.space_before = Pt(18)

                run = p.add_run(f"指标：{label}")

                set_run_font(run, size_pt=9, bold=True)


                fig_base = create_comparison_plot(ind, cars_res, car_folders)


                if fig_base:

                    fig_base.patch.set_facecolor('white')

                    save_fig_to_word(doc, fig_base)

                    plt.close(fig_base)

                # else:

                    # p_no_fig = doc.add_paragraph("[该指标无绘图数据]")

                    # set_run_font(p_no_fig.runs[0], size_pt=8, italic=True, color=RGBColor(150, 150, 150))


                # 5. 生成专项分析图

                for car in car_folders:

                    if label in cars_res[car]['figs']:

                        fig_data = cars_res[car]['figs'][label]

                        # 确保 fig_data 不是 None 或空列表

                        if not fig_data: continue


                        p_cap = doc.add_paragraph(f"[{car}] 分析特征图:")

                        p_cap.paragraph_format.space_before = Pt(6)

                        set_run_font(p_cap.runs[0], size_pt=8)

                        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


                        if isinstance(fig_data, list) and len(fig_data) > 0:

                            # 先过滤掉列表里可能为空(None)的对象，获取真实的图片数量

                            valid_figs = [f for f in fig_data if f]

                            num_figs = len(valid_figs)


                            if num_figs > 0:

                                # 核心逻辑：最多 2 列，计算需要的总行数

                                cols = min(num_figs, 2)

                                rows = math.ceil(num_figs / 2)


                                # 创建多行多列的表格

                                table_img = doc.add_table(rows=rows, cols=cols)

                                table_img.autofit = False


                                # 将图片按顺序填入网格

                                for idx, fig_obj in enumerate(valid_figs):

                                    row_idx = idx // 2  # 整除 2 决定在第几行 (0, 0, 1, 1, 2, 2...)

                                    col_idx = idx % 2  # 对 2 取余决定在第几列 (0, 1, 0, 1, 0, 1...)


                                    cell = table_img.cell(row_idx, col_idx)

                                    save_fig_to_cell(cell, fig_obj)

                        else:

                            save_fig_to_word(doc, fig_data)


                # 6. 【新建表格】因为有图，所以开启一个新的数据表

                active_table = doc.add_table(rows=1, cols=len(car_folders) + 2)

                active_table.style = 'Table Grid'

                set_table_layout(active_table, col_widths_list)


                row0 = active_table.rows[0].cells

                row0[0].text = "对比项"

                if row0[0].paragraphs[0].runs:

                    set_run_font(row0[0].paragraphs[0].runs[0], bold=True, size_pt=header_font_size)

                set_cell_border(row0[0], top=border_heavy, bottom=border_thin)

                row0[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


                row0[1].text = "计算方案"

                if row0[1].paragraphs[0].runs:

                    set_run_font(row0[1].paragraphs[0].runs[0], bold=True, size_pt=header_font_size)

                set_cell_border(row0[1], top=border_heavy, bottom=border_thin)

                row0[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


                for i, car in enumerate(car_folders):

                    cell_h = row0[i + 2]

                    cell_h.text = car

                    if cell_h.paragraphs[0].runs:

                        set_run_font(cell_h.paragraphs[0].runs[0], bold=True, size_pt=header_font_size)

                    cell_h.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                    set_cell_border(cell_h, top=border_heavy, bottom=border_thin)


            # =======================================================

            # 分支 B: 如果无图 (has_fig=False)

            # =======================================================

            else:

                if active_table is None:

                    active_table = doc.add_table(rows=1, cols=len(car_folders) + 2)

                    active_table.style = 'Table Grid'

                    set_table_layout(active_table, col_widths_list)

                    row0 = active_table.rows[0].cells


                    p = row0[0].paragraphs[0]

                    p.clear()

                    r = p.add_run("对比项")

                    set_run_font(r, bold=True, size_pt=header_font_size)

                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    set_cell_border(row0[0], top=border_heavy, bottom=border_thin)


                    p_m = row0[1].paragraphs[0]

                    p_m.clear()

                    r_m = p_m.add_run("计算方案")

                    set_run_font(r_m, bold=True, size_pt=header_font_size)

                    p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    set_cell_border(row0[1], top=border_heavy, bottom=border_thin)


                    for i, car in enumerate(car_folders):

                        cell = row0[i + 2]

                        p = cell.paragraphs[0]

                        p.clear()

                        r = p.add_run(car)

                        set_run_font(r, bold=True, size_pt=header_font_size)

                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        set_cell_border(cell, top=border_heavy, bottom=border_thin)


            # =======================================================

            # 通用步骤: 向当前 active_table 添加数据行

            # =======================================================

            new_row = active_table.add_row()

            cells = new_row.cells


            # 第一列：指标名称

            cells[0].text = label

            if cells[0].paragraphs[0].runs:

                set_run_font(cells[0].paragraphs[0].runs[0], bold=True, size_pt=header_font_size)

            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


            # 第二列：写入计算方案 (多行)

            p_method = cells[1].paragraphs[0]

            p_method.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for i, txt in enumerate(method_list):

                run = p_method.add_run(f"{txt}")

                set_run_font(run, size_pt=header_font_size)

                if i < len(method_list) - 1:

                    run.add_break()

            cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            set_cell_border(cells[1])


            # 后续列：写入车型数值并保证行数对齐

            for c_idx, car in enumerate(car_folders):

                val_data = cars_res[car]['metrics'].get(label, "-")

                val_list = val_data if isinstance(val_data, list) else [val_data]


                # 获取行数最大值，保证数值和左侧计算方案的换行对齐

                max_lines = max(len(method_list), len(val_list))


                cell_v = cells[c_idx + 2]  # 列索引从2开始

                p_val = cell_v.paragraphs[0]

                p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER


                for i in range(max_lines):

                    text_val = str(val_list[i]) if i < len(val_list) else "-"

                    run = p_val.add_run(text_val)

                    set_run_font(run, size_pt=9)

                    if i < max_lines - 1:

                        run.add_break()


                cell_v.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                set_cell_border(cell_v)

                plt.close('all')


        doc.add_paragraph()  # 工况结束空一行



# ==========================================

# 3. 汇总表生成逻辑 (触发第一次计算)

# ==========================================

def create_summary_table(doc, export_path, car_folders, calc, configs, pre_calc_data=None):

    print("\n---------- 阶段 1: 生成汇总数据表 ----------")

    # 设置表格宽度

    headers = ["工况", "指标", "计算方法"] + car_folders

    table = doc.add_table(rows=1, cols=len(headers))

    table.style = 'Table Grid'


    table.autofit = False

    table.allow_autofit = False


    # 定义前三列的固定宽度 (单位: cm)

    width_map = {

        0: 1.0,  # 工况列

        1: IND_WIDTH,  # 指标列

        2: CAL_WIDTH  # 计算方法列

    }


    fixed_width_total = sum(width_map.values())

    remaining_width = PAGE_USABLE_WIDTH - fixed_width_total

    num_cars = len(car_folders)


    if num_cars > 0:

        car_col_width = remaining_width / num_cars

    else:

        car_col_width = 0


    for col_idx, column in enumerate(table.columns):

        if col_idx in width_map:

            column.width = Cm(width_map[col_idx])

        else:

            column.width = Cm(car_col_width)


    # 设置表头

    hdr_cells = table.rows[0].cells

    border_heavy = {"sz": 18, "val": "single", "color": "000000"}

    border_thin = {"sz": 4, "val": "single", "color": "000000"}

    header_font_size = 6 if len(car_folders) > 4 else 8


    for i, h in enumerate(headers):

        cell = hdr_cells[i]

        if i in width_map:

            cell.width = Cm(width_map[i])

        else:

            cell.width = Cm(car_col_width)


        p = cell.paragraphs[0]

        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


        run = p.add_run(h)

        set_run_font(run, size_pt=header_font_size, bold=True)

        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        set_cell_border(cell, top=border_heavy, bottom=border_thin)


    for kw, cfg in configs.items():

        if kw == 'common_config': continue


        # 预加载数据并计算

        for car in car_folders:

            car_folder_path = os.path.join(export_path, car)

            path = find_file_fuzzy(os.path.join(car_folder_path, car), kw)

            # 使用标准化读取函数处理别名

            df = safe_read_csv(path) if path else None

            calc.get_condition_results(car, kw, df, configs, car_folder_path)


        indicators = cfg.get('indicators', [])

        if not indicators: continue  # 防止空指标导致报错


        start_row_idx = len(table.rows)  # 记录起始行


        # --- 3. 遍历指标并填充行 ---

        for idx, ind in enumerate(indicators):

            row_cells = table.add_row().cells


            # (Column 0) 工况标题 (只在第一行显示)

            if idx == 0:

                title = cfg.get('title', kw)

                p = row_cells[0].paragraphs[0]

                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                run = p.add_run(title)

                set_run_font(run, size_pt=8, italic=True)


            # (Column 1) 指标名称

            cell_ind = row_cells[1]  # 为了方便操作，先获取单元格对象

            p = cell_ind.paragraphs[0]

            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = p.add_run(ind['label'])

            set_run_font(run, size_pt=8)

            cell_ind.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


            # (Column 2) 计算方案 - 支持多行

            cell_method = row_cells[2]

            p_method = cell_method.paragraphs[0]

            p_method.alignment = WD_ALIGN_PARAGRAPH.CENTER


            raw_method = ind.get('method', '-')

            # 统一转为列表

            method_list = []

            if isinstance(raw_method, list):

                method_list = [str(m) for m in raw_method]

            elif isinstance(raw_method, str):

                method_list = raw_method.split('\n')

            else:

                method_list = [str(raw_method)]


            # 写入方案描述

            for i, txt in enumerate(method_list):

                run = p_method.add_run(f"{txt}")

                set_run_font(run, size_pt=8)

                if i < len(method_list) - 1:

                    run.add_break()  # 换行


            cell_method.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


            # (Column 3+) 车型数值 - 核心修改：支持多值对齐

            for c_idx, car in enumerate(car_folders):

                cell_val = row_cells[c_idx + 3]

                p_val = cell_val.paragraphs[0]

                p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER


                # [获取数据的核心逻辑]

                res = None

                # A. 尝试从缓存获取

                if pre_calc_data and kw in pre_calc_data and car in pre_calc_data[kw]:

                    res = pre_calc_data[kw][car]


                # B. 如果缓存没有，现场算

                if not res:

                    car_folder_path = os.path.join(export_path, car)

                    path = find_file_fuzzy(car_folder_path, kw)

                    df = safe_read_csv(path) if path else None

                    res = calc.get_condition_results(car, kw, df, configs, car_folder_path)


                val_data = res['metrics'].get(ind['label'], "-")


                # 统一转为列表处理

                val_list = []

                if isinstance(val_data, list):

                    val_list = val_data

                else:

                    val_list = [val_data]


                max_lines = max(len(method_list), len(val_list))


                for i in range(max_lines):

                    # 取出对应行的数值

                    if i < len(val_list):

                        text_val = str(val_list[i])

                    else:

                        text_val = "-"  # 这一行没有值


                    run = p_val.add_run(text_val)

                    set_run_font(run, size_pt=8)


                    # 只要不是最后一行，就加换行符，确保和左边的方案列对齐

                    if i < max_lines - 1:

                        run.add_break()


                cell_val.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


        # --- 执行单元格合并与垂直居中 ---

        end_row_idx = len(table.rows) - 1

        if end_row_idx > start_row_idx:

            cell_start = table.cell(start_row_idx, 0)

            cell_end = table.cell(end_row_idx, 0)

            cell_start.merge(cell_end)

            # 设置垂直居中

            cell_start.vertical_alignment = WD_ALIGN_VERTICAL.CENTER



def generate_report(result_folder, configs, carsim_controller=None, pre_calc_data=None):

    doc = Document()

    # --- 1. 设置 A4 纵向 (Portrait) ---

    section = doc.sections[0]

    section.page_width = Cm(PAGE_WIDTH)  # A4 宽

    section.page_height = Cm(29.7)  # A4 高

    section.left_margin = Cm(PAGE_MARGIN)  # 窄边距

    section.right_margin = Cm(PAGE_MARGIN)

    section.top_margin = Cm(PAGE_MARGIN)

    section.bottom_margin = Cm(PAGE_MARGIN)


    # --- 2. 封面与介绍 ---

    heading = doc.add_heading('', 0)

    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = heading.add_run('操稳离线仿真对比分析报告 V260415')

    set_run_font(run, size_pt=15, bold=True)


    p_info = doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_run_font(p_info.runs[0], size_pt=10, italic=True)


    calc = IndicatorCalculator(carsim_controller)

    car_folders = [d for d in os.listdir(result_folder) if os.path.isdir(os.path.join(result_folder, d))]

    car_folders.sort()  # 排序，保证顺序稳定


    # 介绍段落

    p_intro = doc.add_paragraph()

    intro_txt = f"本报告汇总了 {len(car_folders)} 款车型的仿真结果。首先展示所有指标的数值汇总，随后按工况展示详细的数据曲线与具体指标得分。工况道路宽度3.75米。若未读取到传动比、载荷等信息，默认为1，将影响传动比、侧偏柔度等指标结果，传动比计算优先级是：config.txt > CSV Steer_SW/Steer_L1 均值 > 默认值 1.0。, 如输出无上述两列数值，可在仿真结果文件夹车辆的config.txt中修改，再次点击计算。"

    run_intro = p_intro.add_run(intro_txt)

    set_run_font(run_intro, size_pt=11)


    # ================== 读取并写入配置说明 ==================

    mapping_file = os.path.join(result_folder, "model_info.txt")

    if os.path.exists(mapping_file):

        try:

            with open(mapping_file, 'r', encoding='utf-8') as f:

                mapping_text = f.read()


            p_mapping = doc.add_paragraph()

            run_mapping = p_mapping.add_run(mapping_text)


            # 使用更深一点的灰色和稍微小一点的字号 (比如9号) 来展示配置列表

            set_run_font(run_mapping, size_pt=9, color=RGBColor(80, 80, 80))

            p_mapping.paragraph_format.space_after = Pt(18)  # 在这里加上下间距

        except Exception as e:

            print(f"读取配置说明失败: {e}")

    else:

        p_intro.paragraph_format.space_after = Pt(18)  # 兜底，如果没有文件就维持原来的间距

    # ===============================================================


    # 步骤 1: 汇总表 (优先使用 pre_calc_data)

    create_summary_table(doc, result_folder, car_folders, calc, configs, pre_calc_data)


    # 步骤 2: 详情页 (优先使用 pre_calc_data)

    create_detailed_pages(doc, car_folders, calc, configs, pre_calc_data)


    output_filename = f'仿真分析报告_{datetime.now().strftime("%Y%m%d_%H%M")}.docx'

    full_save_path = os.path.join(result_folder, output_filename)

    try:

        doc.save(full_save_path)

        print(f"\n✅ [成功] 报告生成完毕: {full_save_path}")

        return full_save_path

    except Exception as e:

        print(f"\n❌ [保存失败] 文件可能被占用: {e}")

        return None



if __name__ == "__main__":


    start_time = datetime.now()

    print(f"[{start_time.strftime('%H:%M:%S')}] 任务启动...")


    need_simulation = False

    cc = ControlCarsim()


    if need_simulation is True:


        export_path = r'C:\Users\93446\Desktop\test'

        veh_list = [

            ("MYHF_A_20260127_2104 #ORI", "MYHF"),

            ("MYHF_A_20260127_2104 #FruanRying", "MYHF"),

            ("MYHF_A_20260127_2104 #FyingRruan", "MYHF"),

        ]


        cc.set_vehicle_list(veh_list)


        time_path = cc.run_carsim_batch(export_path)

        cc.recover_dataset()


    else:

        time_path = r'C:\Users\93446\Desktop\test\20260205_105707'


    output_file_name = generate_report(time_path, CONFIGS, cc)


    end_time = datetime.now()

    duration = (end_time - start_time).seconds

    print(f"[{end_time.strftime('%H:%M:%S')}] 任务完成，耗时 {duration} 秒。")


    try:

        os.startfile(output_file_name)


    except Exception as e:

        print(f"❌ 无法打开文件: {e}")


