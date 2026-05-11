

import colorsys

import json

import os

import textwrap

import threading

from typing import Dict, Iterable, Optional, Sequence, Set, Tuple

import numpy as np

import pandas as pd

from PyQt5.QtGui import QFontDatabase, QKeySequence, QColor


from typing import Optional, List


from docx.enum.text import WD_ALIGN_PARAGRAPH

from docx.shared import Cm

from matplotlib import pyplot as plt

from openpyxl.drawing.spreadsheet_drawing import AnchorMarker, OneCellAnchor

from openpyxl.drawing.xdr import XDRPositiveSize2D

import io

from openpyxl.drawing.image import Image as XLImage

from openpyxl.utils import coordinate_to_tuple


# --------------------DIL数据处理-----------------------

CONFIG: Dict = {

    "unit_map": {

        "AZ": "m/s²", "AVX": "deg/s", "AVY": "deg/s", "AVZ": "deg/s",

        "PITCH": "deg", "YAW": "deg", "ROLL": "deg", "AX": "m/s²", "AY": "m/s²", "Z": 'm'

    },

    "color_map": {

        "MoogData": plt.cm.Blues,

        "CarsimData": plt.cm.Oranges,

    },

    "colors": [

        '#8B0000',  # 深红

        '#006400',  # 深绿

        '#00008B',  # 深蓝

        '#8B008B',  # 深紫

        '#FF8C00',  # 深橙

        '#8B4513',  # 深棕

        '#2F4F4F',  # 深灰蓝

        '#008080',  # 深青

        '#556B2F',  # 深橄榄绿

        '#9932CC',  # 深兰花紫

        '#8B0000',  # 深红

        '#8FBC8F',  # 深海绿

        '#483D8B',  # 深板岩蓝

        '#2E8B57',  # 深海绿

        '#8B0000',  # 深红

        '#8B008B',  # 深紫

    ],

    "font_styles": {

        "title_fontsize": 16,

        "subtitle_fontsize": 10,

        "axis_fontsize": 10,

        "legend_fontsize": 8,

        "stats_fontsize": 9,

        "text_fontsize": 8,

        "ptp_fontsize": 9

    },

    "axis_labels": {"time": "时间 (秒)"},

    "data_processing_rules": {

        "MoogData": {

            "Pitch": {"unit_conversion": "rad_to_deg", "invert": True},

            "Yaw": {"unit_conversion": "rad_to_deg", "invert": True},

            "Roll": {"unit_conversion": "rad_to_deg"},

            "Avx": {"unit_conversion": "rad_to_deg"},

            "Avy": {"unit_conversion": "rad_to_deg", "invert": True},

            "Avz": {"unit_conversion": "rad_to_deg", "invert": True},

            "Ay": {"invert": True},

            "Az": {"invert": True},

        },

        "CarsimData": {

            "Az": {"scale_factor": 9.81},

            "Ay": {"scale_factor": 9.81},

            "Ax": {"scale_factor": 9.81}

        },

        "marker_styles": {"peak": "o", "valley": "v", "sec_peak": "o"},

        "marker_sizes": {"peak": 30, "valley": 30, "sec_peak": 30},

    }


}


# 定义 Word 中图片的显示宽度 (单位: 厘米)

WORD_IMG_WIDTH = 7


_DEF_ENCODINGS: Sequence[str] = ("utf-8", "gbk", "gb2312", "latin1")



class ConfigProcessor:

    def __init__(self, config: Optional[Dict] = None) -> None:

        self.config = (config or CONFIG).copy()

        self.data_processing_rules: Dict = self.config.get("data_processing_rules", {})

        self.color_map: Dict = self.config.get("color_map", {}).copy()

        self.colors = self.config.get("colors", {}).copy()

        self.unit_map: Dict = self.config.get("unit_map", {}).copy()

        self.axis_labels: Dict = self.config.get("axis_labels", {})

        styles = self.config.get("font_styles", {})

        for k, v in styles.items():

            setattr(self, k, v)

        self.title_fontsize = getattr(self, "title_fontsize", 16)

        self.subtitle_fontsize = getattr(self, "subtitle_fontsize", 10)

        self.axis_fontsize = getattr(self, "axis_fontsize", 10)

        self.legend_fontsize = getattr(self, "legend_fontsize", 8)

        self.stats_fontsize = getattr(self, "stats_fontsize", 9)

        self.text_fontsize = getattr(self, "text_fontsize", 8)

        self.ptp_fontsize = getattr(self, "ptp_fontsize", 9)

        self._init_chinese_font()


    def _init_chinese_font(self) -> None:

        font_db = QFontDatabase()

        fams = set(font_db.families())

        for name in ("Microsoft YaHei", "SimHei", "SimSun", "KaiTi", "FangSong"):

            if name in fams:

                plt.rcParams['font.sans-serif'] = [name]

                plt.rcParams['axes.unicode_minus'] = False

                return


    def get_palette(self, n, source):

        # 如果是overlay模式，使用完全不同的色系

        if source == "overlay":

            # 确保颜色数量足够

            if n <= len(self.colors):

                return self.colors[:n]


            # 如果需要更多颜色，使用HSL生成补充颜色

            colors = self.colors.copy()

            for i in range(len(self.colors), n):

                # 在色相环上均匀分布

                hue = i / n

                # 保持高饱和度

                saturation = 0.9

                # 设置低亮度（深色）

                lightness = 0.3

                # 转换为RGB

                r, g, b = colorsys.hls_to_rgb(hue, lightness, saturation)

                # 转换为十六进制颜色

                hex_color = "#{:02x}{:02x}{:02x}".format(

                    int(r * 255), int(g * 255), int(b * 255)

                )

                colors.append(hex_color)

            return colors

        else:

            cmap = self.color_map.get(source, plt.cm.tab10)

            return cmap(np.linspace(0.35, 0.95, n))



def safe_idx(name, cols):

    return cols.index(name) if name in cols else 0



def insert_excel_image(ws, fig, anchor, cell_w, cell_h):

    """

    将 Matplotlib Figure 居中插入 Excel

    """

    img_buffer = io.BytesIO()

    fig.savefig(img_buffer, format='png', dpi=200, bbox_inches='tight', facecolor='white')

    img_buffer.seek(0)

    img = XLImage(img_buffer)


    max_w = cell_w - 10

    max_h = cell_h - 10


    img_w = img.width

    img_h = img.height


    # 计算缩放因子 (Fit Inside)

    scale = min(max_w / img_w, max_h / img_h)


    # 应用缩放

    new_w = int(img_w * scale)

    new_h = int(img_h * scale)

    img.width = new_w

    img.height = new_h


    EMU = 9525


    # (容器宽 - 图片宽) / 2

    offset_x = int((cell_w - new_w) / 2 * EMU)

    offset_y = int((cell_h - new_h) / 2 * EMU)


    row, col = coordinate_to_tuple(anchor)

    marker = AnchorMarker(col=col - 1, colOff=offset_x, row=row - 1, rowOff=offset_y)


    size = XDRPositiveSize2D(cx=new_w * EMU, cy=new_h * EMU)


    img.anchor = OneCellAnchor(_from=marker, ext=size)

    ws.add_image(img)



def load_configs(path):

    """加载 JSON 配置"""

    if not os.path.exists(path):

        print(f"❌ [错误] 配置文件不存在: {path}")

        return {}

    with open(path, 'r', encoding='utf-8') as f:

        data = json.load(f)

        return data



def configure_matplotlib_style():

    """配置 Matplotlib 全局样式：字体、字号、线宽等"""


    plt.rcParams['font.family'] = 'serif'

    plt.rcParams['font.serif'] = ['SimSun', 'Times New Roman']

    plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示为方框的问题

    plt.rcParams['figure.max_open_warning'] = 50


    plt.rcParams.update({

        'font.size': 11,  # 默认全局字号

        'axes.titlesize': 11,  # 子图标题 (Title)

        'axes.labelsize': 11,  # 轴标签 (Label: Ay [g])

        'xtick.labelsize': 11,  # X轴刻度 (Ticks)

        'ytick.labelsize': 11,  # Y轴刻度

        'legend.fontsize': 9,  # 图例 (Legend)

        'figure.titlesize': 12  # 整个画布的大标题

    })


    # 3. 线条与网格样式

    plt.rcParams.update({

        'lines.linewidth': 1.5,  # 默认线宽

        'lines.markersize': 6,  # 默认点的大小

        'axes.grid': True,  # 默认开启网格

        'grid.linestyle': ':',  # 网格线型 (虚线)

        'grid.alpha': 0.6,  # 网格透明度

        'axes.linewidth': 1.0,  # 坐标轴边框粗细

    })


    # 4. 布局与保存

    plt.rcParams.update({

        'figure.figsize': (7, 4),


        'figure.constrained_layout.use': False,

        'figure.autolayout': False,


        'savefig.dpi': 200,

        'savefig.bbox': 'tight',

        'savefig.pad_inches': 0.1

    })



def make_legend_outside(ax=None, wrap_len=28):

    """

    :param wrap_len: 每行强制换行的字符数 (例如 10)

    """

    if ax is None:

        ax = plt.gca()


    handles, labels = ax.get_legend_handles_labels()


    if handles:

        final_labels = []

        for lbl in labels:

            lbl_str = str(lbl)

            new_lbl = '\n'.join([lbl_str[i:i+wrap_len] for i in range(0, len(lbl_str), wrap_len)])

            final_labels.append(new_lbl)


        ax.legend(handles, final_labels,

                  bbox_to_anchor=(1.01, 1),

                  loc='upper left',

                  borderaxespad=0)


    fig = ax.get_figure()

    # fig.tight_layout(pad=1.0)

    # fig.subplots_adjust(right=0.74) # 预留空间给变宽的图例

    fig.set_layout_engine('constrained')


def apply_global_styles(widget):

    """

    全局样式应用函数

    :param widget: 需要应用样式的窗口或组件对象 (传入 self)

    """

    widget.setStyleSheet("""

        /* 1. 全局背景与字体 */

        QWidget { 

            background-color: white; 

            font-family: "Times New Roman", "SimSun"; 

            font-size: 10pt; 

        }


        /* 2. GroupBox 样式 */

        QGroupBox { 

            font-weight: bold; 

            border: 1px solid #ccc; 

            border-radius: 5px; 

            margin-top: 10px; 

        }

        QGroupBox::title { 

            subcontrol-origin: margin; 

            subcontrol-position: top left; 

            padding: 0 5px; 

        }


        /* 3. 按钮样式 */

        QPushButton { 

                        background-color: #f0f0f0; 

                        border: 1px solid #ccc; 

                        border-radius: 4px; 

                        padding: 5px 10px;

                        min-width: 60px;

                    }

        QPushButton:hover { border: 2px solid #138496;

                            color: white; 

                            background-color: #138496; }

        QPushButton:checked { 

                                background-color: #138496; 

                                color: white; 

                                border: 2px solid #138496; 

                                font-weight: bold;}

        QPushButton:pressed {

            background-color: #d4d4d4;

        }

                          

        /* 4. 滚动条美化 */

        QScrollBar:vertical {

            border: none;

            background: #f0f0f0;

            width: 10px;

            margin: 0px;

        }

        QScrollBar::handle:vertical {

            background: #cdcdcd;

            min-height: 20px;

            border-radius: 5px;

        }

        QScrollBar:horizontal {

            border: none;

            background: #f0f0f0;

            height: 10px;

        }

        QScrollBar::handle:horizontal {

            background: #cdcdcd;

            min-width: 20px;

            border-radius: 5px;

        }

        

        QProgressBar {

            border: 1px solid #ccc;     /* 边框颜色 */

            border-radius: 5px;         /* 圆角 */

            text-align: center;         /* 文字居中 */

            background-color: #f0f0f0;  /* 轨道背景色 (浅灰) */

            color: #000000;                /* 文字颜色 */

        }

        QProgressBar::chunk {

            background-color: #138496;  /* 进度条填充色 (与按钮一致的青色) */

            border-radius: 3px;         /* 填充部分的圆角，稍微小一点以适应容器 */

            margin: 1px;                /* 留出一点空隙，产生"内嵌"感 */

        }

    """)



def save_fig_to_word(doc, fig):

    buf = io.BytesIO()

    fig.savefig(buf, format='png', bbox_inches='tight', dpi=150)

    buf.seek(0)

    p = doc.add_paragraph()

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p.add_run().add_picture(buf, width=Cm(WORD_IMG_WIDTH))

    plt.close(fig)



def save_fig_to_cell(cell, fig):

    """将 Matplotlib 图像保存并插入到指定的 Word 单元格中"""

    buf = io.BytesIO()

    fig.savefig(buf, format='png', bbox_inches='tight', dpi=150)

    buf.seek(0)


    # 获取单元格里的第一个段落，设置居中

    p = cell.paragraphs[0]

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # 插入图片

    run = p.add_run()

    run.add_picture(buf, width=Cm(WORD_IMG_WIDTH))


    # 记得关闭图释放内存

    plt.close(fig)







